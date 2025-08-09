import json
import os
import sys
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QGridLayout, QTabWidget, QLabel, 
                             QLineEdit, QPushButton, QTableWidget, QTableWidgetItem,
                             QComboBox, QTextEdit, QGroupBox, QMessageBox, 
                             QFileDialog, QDialog, QScrollArea, QFormLayout,
                             QHeaderView, QFrame, QDateEdit, QSpinBox, QDoubleSpinBox,
                             QCheckBox, QStackedWidget, QSizePolicy, QDialogButtonBox)
from PyQt5.QtCore import Qt, QDate, pyqtSignal, QSize
from PyQt5.QtGui import QFont, QIcon, QPalette, QColor, QTextCursor
import datetime
import pandas as pd
from tabulate import tabulate
import hashlib
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import numpy as np
from fpdf import FPDF
import subprocess
import tempfile

# تحسين الخطوط وتكبيرها
LARGE_FONT = QFont("Arial", 12)
MEDIUM_FONT = QFont("Arial", 11)
SMALL_FONT = QFont("Arial", 10)

class LoginDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("تسجيل الدخول")
        self.setFixedSize(450, 300)
        self.setModal(True)
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # Title
        title = QLabel("نظام إدارة الموظفين - Perfection v3")
        title.setAlignment(Qt.AlignCenter)
        title_font = QFont("Arial", 16, QFont.Bold)
        title.setFont(title_font)
        title.setStyleSheet("color: #2c3e50; margin: 20px;")
        layout.addWidget(title)
        
        # Form layout
        form_layout = QFormLayout()
        form_layout.setLabelAlignment(Qt.AlignRight)
        
        self.username_input = QLineEdit()
        self.username_input.setPlaceholderText("أدخل اسم المستخدم")
        self.username_input.setFont(LARGE_FONT)
        form_layout.addRow("اسم المستخدم:", self.username_input)
        
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.Password)
        self.password_input.setPlaceholderText("أدخل كلمة المرور")
        self.password_input.setFont(LARGE_FONT)
        form_layout.addRow("كلمة المرور:", self.password_input)
        
        layout.addLayout(form_layout)
        
        # Login button
        self.login_btn = QPushButton("دخول")
        self.login_btn.setFont(LARGE_FONT)
        self.login_btn.clicked.connect(self.accept)
        layout.addWidget(self.login_btn)
        
        # Connect Enter key to login
        self.password_input.returnPressed.connect(self.accept)
        
        self.setLayout(layout)
        
    def get_credentials(self):
        return self.username_input.text(), self.password_input.text()

class EmployeeEditDialog(QDialog):
    def __init__(self, employee_data, employee_type, parent=None):
        super().__init__(parent)
        self.setWindowTitle("تعديل بيانات الموظف")
        self.setFixedSize(600, 400)
        self.employee_data = employee_data
        self.employee_type = employee_type
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # Employee Type
        type_label = QLabel(f"نوع الموظف: {'بحصص' if self.employee_type == 'session' else 'راتب ثابت'}")
        type_label.setFont(LARGE_FONT)
        type_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(type_label)
        
        # Form layout
        form_layout = QFormLayout()
        form_layout.setLabelAlignment(Qt.AlignRight)
        
        self.name_input = QLineEdit()
        self.name_input.setFont(LARGE_FONT)
        self.name_input.setText(self.employee_data.get('name', ''))
        form_layout.addRow("الاسم:", self.name_input)
        
        self.phone_input = QLineEdit()
        self.phone_input.setFont(LARGE_FONT)
        self.phone_input.setText(self.employee_data.get('phone', ''))
        form_layout.addRow("رقم الهاتف:", self.phone_input)
        
        if self.employee_type == 'session':
            self.rate_input = QDoubleSpinBox()
            self.rate_input.setFont(LARGE_FONT)
            self.rate_input.setMaximum(9999.99)
            self.rate_input.setDecimals(2)
            self.rate_input.setValue(self.employee_data.get('current_rate', 0))
            form_layout.addRow("سعر الحصة الأساسي:", self.rate_input)
        else:
            self.salary_input = QDoubleSpinBox()
            self.salary_input.setFont(LARGE_FONT)
            self.salary_input.setMaximum(99999.99)
            self.salary_input.setDecimals(2)
            self.salary_input.setValue(self.employee_data.get('monthly_salary', 0))
            form_layout.addRow("الراتب الأساسي:", self.salary_input)
        
        layout.addLayout(form_layout)
        
        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        button_box.setFont(LARGE_FONT)
        layout.addWidget(button_box)
        
        self.setLayout(layout)
        
    def get_updated_data(self):
        return {
            'name': self.name_input.text(),
            'phone': self.phone_input.text(),
            'current_rate': self.rate_input.value() if self.employee_type == 'session' else 0,
            'monthly_salary': self.salary_input.value() if self.employee_type == 'fixed' else 0
        }

class ExcelImportDialog(QDialog):
    def __init__(self, title, columns, parent=None):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.setFixedSize(700, 500)
        self.columns = columns
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # File selection
        file_layout = QHBoxLayout()
        self.file_path = QLineEdit()
        self.file_path.setFont(MEDIUM_FONT)
        self.file_path.setReadOnly(True)
        file_layout.addWidget(self.file_path)
        
        browse_btn = QPushButton("استعراض")
        browse_btn.setFont(MEDIUM_FONT)
        browse_btn.clicked.connect(self.browse_file)
        file_layout.addWidget(browse_btn)
        
        layout.addLayout(file_layout)
        
        # Preview table
        self.preview_table = QTableWidget()
        self.preview_table.setFont(MEDIUM_FONT)
        self.preview_table.setColumnCount(len(self.columns))
        self.preview_table.setHorizontalHeaderLabels(self.columns)
        self.preview_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.preview_table)
        
        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        button_box.setFont(MEDIUM_FONT)
        layout.addWidget(button_box)
        
        self.setLayout(layout)
        
    def browse_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "اختر ملف Excel", "", "Excel Files (*.xlsx *.xls)"
        )
        
        if file_path:
            self.file_path.setText(file_path)
            self.load_excel_data(file_path)
    
    def load_excel_data(self, file_path):
        try:
            df = pd.read_excel(file_path)
            self.preview_table.setRowCount(len(df))
            
            for i, row in df.iterrows():
                for j, col in enumerate(self.columns):
                    item = QTableWidgetItem(str(row[j]) if j < len(row) else QTableWidgetItem(""))
                    self.preview_table.setItem(i, j, item)
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"فشل تحميل الملف: {str(e)}")
    
    def get_data(self):
        data = []
        for i in range(self.preview_table.rowCount()):
            row_data = {}
            for j, col in enumerate(self.columns):
                item = self.preview_table.item(i, j)
                row_data[col] = item.text() if item else ""
            data.append(row_data)
        return data

class EnhancedEmployeeSystem(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("نظام إدارة الموظفين - Perfection v3")
        self.setGeometry(100, 100, 1600, 900)
        
        # Initialize data structures
        self.employees = {}
        self.other_employees = {}
        self.users = {"admin": self.hash_password("admin123")}
        self.current_user = None
        self.current_month = datetime.datetime.now().month
        self.current_year = datetime.datetime.now().year
        self.data_file = "employee_data.json"
        
        # Load saved data
        self.load_data()
        
        # Apply styling
        self.apply_styles()
        
        # Show login dialog
        if self.authenticate():
            self.setup_ui()
        else:
            sys.exit()
    
    def apply_styles(self):
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QTabWidget::pane {
                border: 1px solid #c0c0c0;
                background-color: white;
                border-radius: 5px;
            }
            QTabWidget::tab-bar {
                alignment: center;
            }
            QTabBar::tab {
                background-color: #e1e1e1;
                border: 1px solid #c0c0c0;
                padding: 10px 25px;
                margin-right: 2px;
                border-top-left-radius: 5px;
                border-top-right-radius: 5px;
                font-size: 12px;
            }
            QTabBar::tab:selected {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
            }
            QTabBar::tab:hover {
                background-color: #ddd;
            }
            QGroupBox {
                font-weight: bold;
                border: 2px solid #cccccc;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 15px;
                background-color: white;
                font-size: 12px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #2c3e50;
                font-size: 12px;
            }
            QLabel {
                color: #2c3e50;
                font-size: 12px;
            }
            QLineEdit, QComboBox, QSpinBox, QDoubleSpinBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                background-color: white;
                font-size: 12px;
                height: 35px;
            }
            QLineEdit:focus, QComboBox:focus, QSpinBox:focus, QDoubleSpinBox:focus {
                border-color: #4CAF50;
            }
            QPushButton {
                padding: 10px 20px;
                font-weight: bold;
                border: none;
                border-radius: 4px;
                background-color: #4CAF50;
                color: white;
                font-size: 12px;
                height: 40px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:pressed {
                background-color: #3d8b40;
            }
            QPushButton.edit-btn {
                background-color: #2196F3;
            }
            QPushButton.edit-btn:hover {
                background-color: #1976D2;
            }
            QPushButton.delete-btn {
                background-color: #f44336;
            }
            QPushButton.delete-btn:hover {
                background-color: #da190b;
            }
            QPushButton.export-btn {
                background-color: #FF9800;
            }
            QPushButton.export-btn:hover {
                background-color: #F57C00;
            }
            QTableWidget {
                gridline-color: #e0e0e0;
                background-color: white;
                selection-background-color: #0078D7;
                font-size: 12px;
            }
            QTableWidget::item {
                padding: 8px;
                border-bottom: 1px solid #e0e0e0;
            }
            QHeaderView::section {
                background-color: #f0f0f0;
                color: #2c3e50;
                padding: 10px;
                border: 1px solid #c0c0c0;
                font-weight: bold;
                font-size: 12px;
            }
            QTextEdit {
                background-color: white;
                border: 1px solid #ddd;
                border-radius: 4px;
                font-family: 'Arial', sans-serif;
                font-size: 12px;
            }
            QDateEdit {
                height: 35px;
            }
        """)
    
    def hash_password(self, password):
        return hashlib.sha256(password.encode()).hexdigest()
    
    def authenticate(self):
        login_dialog = LoginDialog()
        if login_dialog.exec_() == QDialog.Accepted:
            username, password = login_dialog.get_credentials()
            if username in self.users and self.users[username] == self.hash_password(password):
                self.current_user = username
                return True
            else:
                QMessageBox.critical(self, "خطأ", "اسم المستخدم أو كلمة المرور غير صحيحة")
                return False
        return False
    
    def save_data(self):
        """حفظ جميع البيانات في ملف JSON"""
        data = {
            'employees': self.employees,
            'other_employees': self.other_employees,
            'users': self.users
        }
        try:
            with open(self.data_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"فشل حفظ البيانات: {str(e)}")

    def load_data(self):
        """تحميل البيانات من ملف JSON"""
        if os.path.exists(self.data_file):
            try:
                with open(self.data_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.employees = data.get('employees', {})
                    self.other_employees = data.get('other_employees', {})
                    self.users = data.get('users', {"admin": self.hash_password("admin123")})
            except Exception as e:
                QMessageBox.critical(self, "خطأ", f"فشل تحميل البيانات: {str(e)}")
    
    def setup_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        layout = QVBoxLayout()
        central_widget.setLayout(layout)
        
        # Create tab widget
        self.tab_widget = QTabWidget()
        self.tab_widget.setFont(LARGE_FONT)
        layout.addWidget(self.tab_widget)
        
        # Create tabs
        self.create_employee_tab()
        self.create_attendance_tab()
        self.create_bonus_tab()
        self.create_advance_tab()
        self.create_reports_tab()
        
        # Load initial data
        self.update_employee_lists()
        
        self.show()
    
    def create_employee_tab(self):
        tab = QWidget()
        self.tab_widget.addTab(tab, "إدارة الموظفين")
        
        layout = QGridLayout()
        tab.setLayout(layout)
        
        # Regular Employees Group
        reg_group = QGroupBox("موظفين بحصص")
        reg_layout = QFormLayout()
        reg_layout.setLabelAlignment(Qt.AlignRight)
        
        self.reg_name = QLineEdit()
        self.reg_name.setFont(LARGE_FONT)
        self.reg_name.setPlaceholderText("أدخل اسم الموظف")
        reg_layout.addRow("الاسم:", self.reg_name)
        
        self.reg_phone = QLineEdit()
        self.reg_phone.setFont(LARGE_FONT)
        self.reg_phone.setPlaceholderText("أدخل رقم الهاتف")
        reg_layout.addRow("الهاتف:", self.reg_phone)
        
        self.reg_rate = QDoubleSpinBox()
        self.reg_rate.setFont(LARGE_FONT)
        self.reg_rate.setMaximum(9999.99)
        self.reg_rate.setDecimals(2)
        reg_layout.addRow("سعر الحصة الأساسي:", self.reg_rate)
        
        btn_layout = QHBoxLayout()
        save_btn = QPushButton("حفظ البيانات")
        save_btn.setFont(MEDIUM_FONT)
        save_btn.clicked.connect(self.save_regular_employee)
        btn_layout.addWidget(save_btn)
        
        delete_btn = QPushButton("حذف الموظف")
        delete_btn.setFont(MEDIUM_FONT)
        delete_btn.setProperty("class", "delete-btn")
        delete_btn.clicked.connect(self.delete_regular_employee)
        btn_layout.addWidget(delete_btn)
        
        import_btn = QPushButton("استيراد من Excel")
        import_btn.setFont(MEDIUM_FONT)
        import_btn.setProperty("class", "edit-btn")
        import_btn.clicked.connect(self.import_employees_from_excel)
        btn_layout.addWidget(import_btn)
        
        reg_layout.addRow(btn_layout)
        reg_group.setLayout(reg_layout)
        layout.addWidget(reg_group, 0, 0)
        
        # Other Employees Group
        other_group = QGroupBox("موظفين براتب ثابت")
        other_layout = QFormLayout()
        other_layout.setLabelAlignment(Qt.AlignRight)
        
        self.other_name = QLineEdit()
        self.other_name.setFont(LARGE_FONT)
        self.other_name.setPlaceholderText("أدخل اسم الموظف")
        other_layout.addRow("الاسم:", self.other_name)
        
        self.other_phone = QLineEdit()
        self.other_phone.setFont(LARGE_FONT)
        self.other_phone.setPlaceholderText("أدخل رقم الهاتف")
        other_layout.addRow("الهاتف:", self.other_phone)
        
        self.other_salary = QDoubleSpinBox()
        self.other_salary.setFont(LARGE_FONT)
        self.other_salary.setMaximum(99999.99)
        self.other_salary.setDecimals(2)
        other_layout.addRow("الراتب الأساسي:", self.other_salary)
        
        self.other_month = QComboBox()
        self.other_month.setFont(LARGE_FONT)
        self.other_month.addItems([str(i) for i in range(1, 13)])
        self.other_month.setCurrentIndex(self.current_month - 1)
        other_layout.addRow("الشهر:", self.other_month)
        
        self.other_year = QComboBox()
        self.other_year.setFont(LARGE_FONT)
        self.other_year.addItems([str(i) for i in range(2020, 2031)])
        self.other_year.setCurrentText(str(datetime.datetime.now().year))
        other_layout.addRow("السنة:", self.other_year)
        
        self.other_monthly_salary = QDoubleSpinBox()
        self.other_monthly_salary.setFont(LARGE_FONT)
        self.other_monthly_salary.setMaximum(99999.99)
        self.other_monthly_salary.setDecimals(2)
        other_layout.addRow("الراتب الشهري:", self.other_monthly_salary)
        
        btn_layout = QHBoxLayout()
        save_btn = QPushButton("حفظ البيانات")
        save_btn.setFont(MEDIUM_FONT)
        save_btn.clicked.connect(self.save_other_employee)
        btn_layout.addWidget(save_btn)
        
        delete_btn = QPushButton("حذف الموظف")
        delete_btn.setFont(MEDIUM_FONT)
        delete_btn.setProperty("class", "delete-btn")
        delete_btn.clicked.connect(self.delete_other_employee)
        btn_layout.addWidget(delete_btn)
        
        update_btn = QPushButton("تحديث الراتب")
        update_btn.setFont(MEDIUM_FONT)
        update_btn.setProperty("class", "edit-btn")
        update_btn.clicked.connect(self.update_salary)
        btn_layout.addWidget(update_btn)
        
        other_layout.addRow(btn_layout)
        other_group.setLayout(other_layout)
        layout.addWidget(other_group, 0, 1)
        
        # Employee List
        list_group = QGroupBox("سجل الموظفين")
        list_layout = QVBoxLayout()
        
        self.employee_table = QTableWidget()
        self.employee_table.setFont(MEDIUM_FONT)
        self.employee_table.setColumnCount(6)
        self.employee_table.setHorizontalHeaderLabels(["الاسم", "الهاتف", "النوع", "سعر الحصة/الراتب", "تعديل", "حذف"])
        self.employee_table.horizontalHeader().setStretchLastSection(False)
        self.employee_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.employee_table.cellClicked.connect(self.handle_employee_cell_click)
        
        list_layout.addWidget(self.employee_table)
        list_group.setLayout(list_layout)
        layout.addWidget(list_group, 1, 0, 1, 2)
    
    def create_attendance_tab(self):
        tab = QWidget()
        self.tab_widget.addTab(tab, "تسجيل الحضور")
        
        layout = QGridLayout()
        tab.setLayout(layout)
        
        # Attendance Entry Group
        entry_group = QGroupBox("تسجيل يومي")
        entry_layout = QFormLayout()
        entry_layout.setLabelAlignment(Qt.AlignRight)
        
        self.att_date = QDateEdit()
        self.att_date.setFont(LARGE_FONT)
        self.att_date.setDate(QDate.currentDate())
        self.att_date.setCalendarPopup(True)
        self.att_date.dateChanged.connect(self.update_day_name)
        entry_layout.addRow("التاريخ:", self.att_date)
        
        self.day_name = QLabel()
        self.day_name.setFont(LARGE_FONT)
        self.update_day_name()
        entry_layout.addRow("اليوم:", self.day_name)
        
        btn_layout = QHBoxLayout()
        
        reset_btn = QPushButton("إعادة تعيين")
        reset_btn.setFont(MEDIUM_FONT)
        reset_btn.setProperty("class", "edit-btn")
        reset_btn.clicked.connect(self.reset_attendance_date)
        btn_layout.addWidget(reset_btn)
        
        daily_btn = QPushButton("تسجيل حضور لليوم")
        daily_btn.setFont(MEDIUM_FONT)
        daily_btn.clicked.connect(self.open_daily_attendance_window)
        btn_layout.addWidget(daily_btn)
        
        edit_btn = QPushButton("تعديل حضور")
        edit_btn.setFont(MEDIUM_FONT)
        edit_btn.setProperty("class", "edit-btn")
        edit_btn.clicked.connect(self.edit_attendance)
        btn_layout.addWidget(edit_btn)
        
        import_btn = QPushButton("استيراد من Excel")
        import_btn.setFont(MEDIUM_FONT)
        import_btn.setProperty("class", "edit-btn")
        import_btn.clicked.connect(self.import_attendance_from_excel)
        btn_layout.addWidget(import_btn)
        
        entry_layout.addRow(btn_layout)
        entry_group.setLayout(entry_layout)
        layout.addWidget(entry_group, 0, 0)
        
        # Filter Group
        filter_group = QGroupBox("تصفية حسب التاريخ")
        filter_layout = QFormLayout()
        filter_layout.setLabelAlignment(Qt.AlignRight)
        
        self.filter_from_date = QDateEdit()
        self.filter_from_date.setFont(LARGE_FONT)
        self.filter_from_date.setCalendarPopup(True)
        self.filter_from_date.setDate(QDate.currentDate().addDays(-7))
        filter_layout.addRow("من تاريخ:", self.filter_from_date)
        
        self.filter_to_date = QDateEdit()
        self.filter_to_date.setFont(LARGE_FONT)
        self.filter_to_date.setCalendarPopup(True)
        self.filter_to_date.setDate(QDate.currentDate())
        filter_layout.addRow("إلى تاريخ:", self.filter_to_date)
        
        filter_btn_layout = QHBoxLayout()
        
        filter_btn = QPushButton("تصفية")
        filter_btn.setFont(MEDIUM_FONT)
        filter_btn.clicked.connect(self.filter_attendance_by_date)
        filter_btn_layout.addWidget(filter_btn)
        
        reset_filter_btn = QPushButton("إعادة تعيين")
        reset_filter_btn.setFont(MEDIUM_FONT)
        reset_filter_btn.setProperty("class", "edit-btn")
        reset_filter_btn.clicked.connect(self.reset_attendance_filter)
        filter_btn_layout.addWidget(reset_filter_btn)
        
        filter_layout.addRow(filter_btn_layout)
        filter_group.setLayout(filter_layout)
        layout.addWidget(filter_group, 0, 1)
        
        # Attendance List
        list_group = QGroupBox("سجل الحضور")
        list_layout = QVBoxLayout()
        
        self.attendance_table = QTableWidget()
        self.attendance_table.setFont(MEDIUM_FONT)
        self.attendance_table.setColumnCount(8)
        self.attendance_table.setHorizontalHeaderLabels([
            "الموظف", "التاريخ", "اليوم", "الحصص", "البونص اليومي", "الهاتف", "تعديل", "حذف"
        ])
        self.attendance_table.horizontalHeader().setStretchLastSection(False)
        self.attendance_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.attendance_table.cellClicked.connect(self.handle_attendance_cell_click)
        
        list_layout.addWidget(self.attendance_table)
        list_group.setLayout(list_layout)
        layout.addWidget(list_group, 1, 0, 1, 2)
    
    # ... (بقية الأكواد للتبويبات الأخرى بنفس النمط) ...
    
    # Employee Management Methods
    def save_regular_employee(self):
        name = self.reg_name.text().strip()
        phone = self.reg_phone.text().strip()
        rate = self.reg_rate.value()
        
        if not name or rate <= 0:
            QMessageBox.warning(self, "خطأ", "الاسم وسعر الحصة مطلوبان")
            return
        
        if name in self.employees or name in self.other_employees:
            QMessageBox.warning(self, "خطأ", "اسم الموظف موجود بالفعل")
            return
        
        self.employees[name] = {
            'phone': phone,
            'current_rate': rate,
            'attendance': {},
            'performance_bonus': {},
            'monthly_bonuses': {},
            'deductions': {},
            'advances': {},
            'advance_due_dates': {},
            'monthly_rates': {}
        }
        
        self.update_employee_lists()
        self.save_data()
        QMessageBox.information(self, "تم", f"تم حفظ بيانات الموظف {name}")
        self.clear_regular_employee_fields()
    
    def save_other_employee(self):
        name = self.other_name.text().strip()
        phone = self.other_phone.text().strip()
        salary = self.other_salary.value()
        
        if not name or salary <= 0:
            QMessageBox.warning(self, "خطأ", "الاسم والراتب الأساسي مطلوبان")
            return
        
        if name in self.employees or name in self.other_employees:
            QMessageBox.warning(self, "خطأ", "اسم الموظف موجود بالفعل")
            return
        
        self.other_employees[name] = {
            'phone': phone,
            'monthly_salary': salary,
            'monthly_salaries': {},
            'monthly_bonuses': {},
            'deductions': {},
            'advances': {},
            'advance_due_dates': {}
        }
        
        self.update_employee_lists()
        self.save_data()
        QMessageBox.information(self, "تم", f"تم حفظ بيانات الموظف {name}")
        self.clear_other_employee_fields()
    
    def import_employees_from_excel(self):
        dialog = ExcelImportDialog("استيراد موظفين من Excel", ["الاسم", "الهاتف", "النوع", "سعر الحصة/الراتب"])
        if dialog.exec_() == QDialog.Accepted:
            employees_data = dialog.get_data()
            
            for emp in employees_data:
                name = emp.get("الاسم", "").strip()
                phone = emp.get("الهاتف", "").strip()
                emp_type = emp.get("النوع", "").strip()
                value_str = emp.get("سعر الحصة/الراتب", "0").strip()
                
                if not name:
                    continue
                
                try:
                    value = float(value_str)
                except ValueError:
                    value = 0
                
                if emp_type == "بحصص":
                    if name not in self.employees:
                        self.employees[name] = {
                            'phone': phone,
                            'current_rate': value,
                            # ... other fields ...
                        }
                elif emp_type == "راتب ثابت":
                    if name not in self.other_employees:
                        self.other_employees[name] = {
                            'phone': phone,
                            'monthly_salary': value,
                            # ... other fields ...
                        }
            
            self.update_employee_lists()
            self.save_data()
            QMessageBox.information(self, "تم", f"تم استيراد {len(employees_data)} موظف")
    
    def import_attendance_from_excel(self):
        dialog = ExcelImportDialog("استيراد حضور من Excel", ["الموظف", "عدد الحصص", "البونص اليومي"])
        if dialog.exec_() == QDialog.Accepted:
            attendance_data = dialog.get_data()
            date_str = self.att_date.date().toString("yyyy-MM-dd")
            
            for record in attendance_data:
                name = record.get("الموظف", "").strip()
                sessions = record.get("عدد الحصص", "0").strip()
                bonus = record.get("البونص اليومي", "0").strip()
                
                if not name or not sessions.isdigit():
                    continue
                
                sessions = int(sessions)
                bonus = float(bonus) if bonus.replace('.', '', 1).isdigit() else 0.0
                
                if name in self.employees:
                    if 'attendance' not in self.employees[name]:
                        self.employees[name]['attendance'] = {}
                    
                    self.employees[name]['attendance'][date_str] = {
                        'sessions': sessions,
                        'daily_bonus': bonus
                    }
            
            self.update_employee_lists()
            self.save_data()
            QMessageBox.information(self, "تم", f"تم استيراد حضور لـ {len(attendance_data)} موظف لليوم {date_str}")
    
    def update_employee_lists(self):
        # Update comboboxes
        all_employees = list(self.employees.keys()) + list(self.other_employees.keys())
        regular_employees = list(self.employees.keys())
        
        # ... (بقية التحديثات) ...
        
        # Update employee table
        self.employee_table.setRowCount(0)
        
        for name, data in self.employees.items():
            row_position = self.employee_table.rowCount()
            self.employee_table.insertRow(row_position)
            
            self.employee_table.setItem(row_position, 0, QTableWidgetItem(name))
            self.employee_table.setItem(row_position, 1, QTableWidgetItem(data.get('phone', '')))
            self.employee_table.setItem(row_position, 2, QTableWidgetItem("بحصص"))
            self.employee_table.setItem(row_position, 3, QTableWidgetItem(f"{data.get('current_rate', 0):.2f}"))
            
            # Edit button
            edit_btn = QPushButton("تعديل")
            edit_btn.setProperty("row", row_position)
            edit_btn.setProperty("class", "edit-btn")
            edit_btn.clicked.connect(self.edit_employee)
            self.employee_table.setCellWidget(row_position, 4, edit_btn)
            
            # Delete button
            delete_btn = QPushButton("حذف")
            delete_btn.setProperty("row", row_position)
            delete_btn.setProperty("class", "delete-btn")
            delete_btn.clicked.connect(self.delete_employee)
            self.employee_table.setCellWidget(row_position, 5, delete_btn)
        
        for name, data in self.other_employees.items():
            row_position = self.employee_table.rowCount()
            self.employee_table.insertRow(row_position)
            
            self.employee_table.setItem(row_position, 0, QTableWidgetItem(name))
            self.employee_table.setItem(row_position, 1, QTableWidgetItem(data.get('phone', '')))
            self.employee_table.setItem(row_position, 2, QTableWidgetItem("راتب ثابت"))
            self.employee_table.setItem(row_position, 3, QTableWidgetItem(f"{data.get('monthly_salary', 0):.2f}"))
            
            # Edit button
            edit_btn = QPushButton("تعديل")
            edit_btn.setProperty("row", row_position)
            edit_btn.setProperty("class", "edit-btn")
            edit_btn.clicked.connect(self.edit_employee)
            self.employee_table.setCellWidget(row_position, 4, edit_btn)
            
            # Delete button
            delete_btn = QPushButton("حذف")
            delete_btn.setProperty("row", row_position)
            delete_btn.setProperty("class", "delete-btn")
            delete_btn.clicked.connect(self.delete_employee)
            self.employee_table.setCellWidget(row_position, 5, delete_btn)
        
        # ... (بقية التحديثات) ...
    
    def edit_employee(self):
        button = self.sender()
        if button:
            row = button.property("row")
            name = self.employee_table.item(row, 0).text()
            emp_type = self.employee_table.item(row, 2).text()
            
            if emp_type == "بحصص" and name in self.employees:
                dialog = EmployeeEditDialog({
                    'name': name,
                    'phone': self.employees[name].get('phone', ''),
                    'current_rate': self.employees[name].get('current_rate', 0)
                }, 'session', self)
                
                if dialog.exec_() == QDialog.Accepted:
                    updated_data = dialog.get_updated_data()
                    # Preserve existing data
                    self.employees[name] = {
                        **self.employees[name],
                        **updated_data
                    }
                    
                    # If name changed, update key
                    if updated_data['name'] != name:
                        self.employees[updated_data['name']] = self.employees.pop(name)
                    
                    self.update_employee_lists()
                    self.save_data()
                    QMessageBox.information(self, "تم", "تم تحديث بيانات الموظف")
            
            elif emp_type == "راتب ثابت" and name in self.other_employees:
                dialog = EmployeeEditDialog({
                    'name': name,
                    'phone': self.other_employees[name].get('phone', ''),
                    'monthly_salary': self.other_employees[name].get('monthly_salary', 0)
                }, 'fixed', self)
                
                if dialog.exec_() == QDialog.Accepted:
                    updated_data = dialog.get_updated_data()
                    # Preserve existing data
                    self.other_employees[name] = {
                        **self.other_employees[name],
                        **updated_data
                    }
                    
                    # If name changed, update key
                    if updated_data['name'] != name:
                        self.other_employees[updated_data['name']] = self.other_employees.pop(name)
                    
                    self.update_employee_lists()
                    self.save_data()
                    QMessageBox.information(self, "تم", "تم تحديث بيانات الموظف")
    
    def delete_employee(self):
        button = self.sender()
        if button:
            row = button.property("row")
            name = self.employee_table.item(row, 0).text()
            emp_type = self.employee_table.item(row, 2).text()
            
            if emp_type == "بحصص" and name in self.employees:
                reply = QMessageBox.question(self, "تأكيد", f"هل أنت متأكد من حذف الموظف {name}؟")
                if reply == QMessageBox.Yes:
                    del self.employees[name]
                    self.update_employee_lists()
                    self.save_data()
                    QMessageBox.information(self, "تم", f"تم حذف الموظف {name}")
            
            elif emp_type == "راتب ثابت" and name in self.other_employees:
                reply = QMessageBox.question(self, "تأكيد", f"هل أنت متأكد من حذف الموظف {name}؟")
                if reply == QMessageBox.Yes:
                    del self.other_employees[name]
                    self.update_employee_lists()
                    self.save_data()
                    QMessageBox.information(self, "تم", f"تم حذف الموظف {name}")
    
    def filter_attendance_by_date(self):
        from_date = self.filter_from_date.date().toPyDate()
        to_date = self.filter_to_date.date().toPyDate()
        
        self.attendance_table.setRowCount(0)
        
        sorted_employees = sorted(self.employees.items(), key=lambda x: x[0])
        
        for name, data in sorted_employees:
            sorted_attendance = sorted(data.get('attendance', {}).items(), key=lambda x: x[0])
            
            for date, record in sorted_attendance:
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d").date()
                
                if date_obj < from_date or date_obj > to_date:
                    continue
                
                day_name = ["الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت", "الأحد"][date_obj.weekday()]
                phone = data.get('phone', '')
                
                row_position = self.attendance_table.rowCount()
                self.attendance_table.insertRow(row_position)
                
                self.attendance_table.setItem(row_position, 0, QTableWidgetItem(name))
                self.attendance_table.setItem(row_position, 1, QTableWidgetItem(date))
                self.attendance_table.setItem(row_position, 2, QTableWidgetItem(day_name))
                self.attendance_table.setItem(row_position, 3, QTableWidgetItem(str(record.get('sessions', 0))))
                self.attendance_table.setItem(row_position, 4, QTableWidgetItem(f"{record.get('daily_bonus', 0):.2f}"))
                self.attendance_table.setItem(row_position, 5, QTableWidgetItem(phone))
                
                # Edit button
                edit_btn = QPushButton("تعديل")
                edit_btn.setProperty("name", name)
                edit_btn.setProperty("date", date)
                edit_btn.setProperty("class", "edit-btn")
                edit_btn.clicked.connect(self.edit_attendance_record)
                self.attendance_table.setCellWidget(row_position, 6, edit_btn)
                
                # Delete button
                delete_btn = QPushButton("حذف")
                delete_btn.setProperty("name", name)
                delete_btn.setProperty("date", date)
                delete_btn.setProperty("class", "delete-btn")
                delete_btn.clicked.connect(self.delete_attendance_record)
                self.attendance_table.setCellWidget(row_position, 7, delete_btn)
    
    def edit_attendance_record(self):
        button = self.sender()
        if button:
            name = button.property("name")
            date = button.property("date")
            
            date_obj = QDate.fromString(date, "yyyy-MM-dd")
            self.att_date.setDate(date_obj)
            self.open_daily_attendance_window(edit_mode=True)
    
    def delete_attendance_record(self):
        button = self.sender()
        if button:
            name = button.property("name")
            date = button.property("date")
            
            if name in self.employees and date in self.employees[name].get('attendance', {}):
                reply = QMessageBox.question(self, "تأكيد", 
                                           f"هل أنت متأكد من حذف تسجيل حضور {name} بتاريخ {date}؟")
                if reply == QMessageBox.Yes:
                    del self.employees[name]['attendance'][date]
                    self.update_employee_lists()
                    self.save_data()
                    QMessageBox.information(self, "تم", "تم حذف تسجيل الحضور")
    
    def export_to_word(self):
        month = int(self.report_month.currentText())
        year = int(self.report_year.currentText())
        
        # Create a new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'تقرير رواتب الموظفين - {month}/{year}', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph(f"تاريخ التقرير: {datetime.date.today().strftime('%Y-%m-%d')}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a table for each employee
        for name, data in {**self.employees, **self.other_employees}.items():
            # Add page break for each employee except the first one
            if doc.paragraphs:
                doc.add_page_break()
            
            # Employee header
            emp_header = doc.add_heading(f"الموظف: {name}", level=1)
            emp_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add employee type and phone
            emp_type = "بحصص" if name in self.employees else "راتب ثابت"
            phone = data.get('phone', '')
            emp_info = doc.add_paragraph(f"النوع: {emp_type} | الهاتف: {phone}")
            emp_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add a table for the report
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set table width to 100% of page
            table.autofit = False
            table.allow_autofit = False
            table.width = Inches(6)
            
            # Add header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'القيمة'
            hdr_cells[1].text = 'البند'
            
            # Add shading to header
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
            hdr_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
            
            # Add employee details to the table
            def add_row(label, value):
                row_cells = table.add_row().cells
                row_cells[0].text = str(value)
                row_cells[1].text = label
                # Set cell width
                row_cells[0].width = Inches(4)
                row_cells[1].width = Inches(2)
            
            # Calculate salary for the month
            report = self.calculate_salary_for_period(
                name, 
                datetime.datetime(year, month, 1), 
                datetime.datetime(year, month, 1) + datetime.timedelta(days=32)
            )
            
            if report:
                if report['type'] == 'بحصص':
                    add_row('سعر الحصة الأساسي', f"{report['base_rate']:.2f}")
                    add_row('سعر الحصة الحالي', f"{report['current_rate']:.2f}")
                    add_row('عدد الحصص العادية', report['sessions'])
                    add_row('راتب الحصص', f"{report['sessions_salary']:.2f}")
                    add_row('عدد حصص الأداء', report['performance_sessions'])
                    add_row('بونص الأداء', f"{report['performance_bonus']:.2f}")
                    add_row('إجمالي البونص اليومي', f"{report['daily_bonus']:.2f}")
                    add_row('البونص الشهري', f"{report['monthly_bonus']:.2f}")
                else:
                    add_row('الراتب الأساسي', f"{report['base_salary']:.2f}")
                    add_row('الراتب الشهري', f"{report['monthly_salary']:.2f}")
                    add_row('البونص الشهري', f"{report['monthly_bonus']:.2f}")
                
                add_row('الخصومات', f"{report['deduction']:.2f}")
                add_row('السلف المستحقة', f"{report['advance']:.2f}")
                
                # Add total salary row with bold font and different background color
                total_row = table.add_row().cells
                total_row[0].text = f"{report['salary']:.2f}"
                total_row[1].text = 'صافي الراتب'
                
                # Apply bold font
                for cell in total_row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Apply background color
                shading_elm = parse_xml(r'<w:shd {} w:fill="B4C6E7"/>'.format(nsdecls('w')))
                total_row[0]._tc.get_or_add_tcPr().append(shading_elm)
                total_row[1]._tc.get_or_add_tcPr().append(shading_elm)
            
            # Add signature
            doc.add_paragraph("\n")
            doc.add_paragraph("توقيع الموظف: ________________")
            doc.add_paragraph("توقيع المدير: ________________")
            doc.add_paragraph(f"تاريخ: {datetime.date.today().strftime('%Y-%m-%d')}")
        
        # Save the document
        file_path, _ = QFileDialog.getSaveFileName(
            self, "حفظ التقرير كملف وورد", f"employee_reports_{month}_{year}.docx",
            "Word files (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                doc.save(file_path)
                QMessageBox.information(self, "تم", f"تم تصدير التقرير إلى {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "خطأ", f"فشل التصدير: {str(e)}")
    
    def send_via_whatsapp(self):
        name = self.report_employee.currentText()
        if not name:
            QMessageBox.warning(self, "خطأ", "الرجاء تحديد الموظف")
            return
        
        # Generate PDF report
        pdf_path = self.generate_pdf_report(name)
        if not pdf_path:
            return
        
        try:
            # This is platform-dependent and may need adjustment
            if sys.platform.startswith('win'):
                # For Windows
                subprocess.Popen(f'start whatsapp://send?text=تقرير الراتب&document={pdf_path}', shell=True)
            elif sys.platform.startswith('darwin'):
                # For macOS
                subprocess.Popen(['open', f'whatsapp://send?text=تقرير الراتب&document={pdf_path}'])
            else:
                # For Linux
                subprocess.Popen(['xdg-open', f'whatsapp://send?text=تقرير الراتب&document={pdf_path}'])
            
            QMessageBox.information(self, "تم", "تم فتح واتساب مع التقرير المرفق")
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"فشل فتح واتساب: {str(e)}")
    
    def generate_pdf_report(self, name):
        month = int(self.report_month.currentText())
        year = int(self.report_year.currentText())
        
        # Calculate salary for the month
        report = self.calculate_salary_for_period(
            name, 
            datetime.datetime(year, month, 1), 
            datetime.datetime(year, month, 1) + datetime.timedelta(days=32)
        )
        
        if not report:
            QMessageBox.warning(self, "خطأ", "لا يوجد بيانات لهذا الموظف")
            return None
        
        # Create PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
        pdf.set_font('DejaVu', '', 14)
        
        # Title
        pdf.cell(0, 10, f"تقرير راتب الموظف - {name}", 0, 1, 'C')
        pdf.ln(5)
        
        # Employee info
        emp_type = "بحصص" if report['type'] == 'بحصص' else "راتب ثابت"
        phone = ""
        if name in self.employees:
            phone = self.employees[name].get('phone', '')
        elif name in self.other_employees:
            phone = self.other_employees[name].get('phone', '')
        
        pdf.cell(0, 10, f"النوع: {emp_type} | الهاتف: {phone}", 0, 1, 'R')
        pdf.cell(0, 10, f"الشهر: {month}/{year}", 0, 1, 'R')
        pdf.cell(0, 10, f"تاريخ التقرير: {datetime.date.today().strftime('%Y-%m-%d')}", 0, 1, 'R')
        pdf.ln(10)
        
        # Report details
        pdf.set_font('DejaVu', '', 12)
        
        if report['type'] == 'بحصص':
            pdf.cell(100, 10, "سعر الحصة الأساسي:", 0, 0, 'R')
            pdf.cell(50, 10, f"{report['base_rate']:.2f}", 0, 1, 'L')
            
            pdf.cell(100, 10, "سعر الحصة الحالي:", 0, 0, 'R')
            pdf.cell(50, 10, f"{report['current_rate']:.2f}", 0, 1, 'L')
            
            pdf.cell(100, 10, "عدد الحصص العادية:", 0, 0, 'R')
            pdf.cell(50, 10, str(report['sessions']), 0, 1, 'L')
            
            pdf.cell(100, 10, "راتب الحصص:", 0, 0, 'R')
            pdf.cell(50, 10, f"{report['sessions_salary']:.2f}", 0, 1, 'L')
            
            pdf.cell(100, 10, "عدد حصص الأداء:", 0, 0, 'R')
            pdf.cell(50, 10, str(report['performance_sessions']), 0, 1, 'L')
            
            pdf.cell(100, 10, "بونص الأداء:", 0, 0, 'R')
            pdf.cell(50, 10, f"{report['performance_bonus']:.2f}", 0, 1, 'L')
            
            pdf.cell(100, 10, "إجمالي البونص اليومي:", 0, 0, 'R')
            pdf.cell(50, 10, f"{report['daily_bonus']:.2f}", 0, 1, 'L')
            
            pdf.cell(100, 10, "البونص الشهري:", 0, 0, 'R')
            pdf.cell(50, 10, f"{report['monthly_bonus']:.2f}", 0, 1, 'L')
        else:
            pdf.cell(100, 10, "الراتب الأساسي:", 0, 0, 'R')
            pdf.cell(50, 10, f"{report['base_salary']:.2f}", 0, 1, 'L')
            
            pdf.cell(100, 10, "الراتب الشهري:", 0, 0, 'R')
            pdf.cell(50, 10, f"{report['monthly_salary']:.2f}", 0, 1, 'L')
            
            pdf.cell(100, 10, "البونص الشهري:", 0, 0, 'R')
            pdf.cell(50, 10, f"{report['monthly_bonus']:.2f}", 0, 1, 'L')
        
        pdf.cell(100, 10, "الخصومات:", 0, 0, 'R')
        pdf.cell(50, 10, f"{report['deduction']:.2f}", 0, 1, 'L')
        
        pdf.cell(100, 10, "السلف المستحقة:", 0, 0, 'R')
        pdf.cell(50, 10, f"{report['advance']:.2f}", 0, 1, 'L')
        
        pdf.ln(10)
        pdf.set_font('DejaVu', 'B', 14)
        pdf.cell(100, 10, "صافي الراتب:", 0, 0, 'R')
        pdf.cell(50, 10, f"{report['salary']:.2f}", 0, 1, 'L')
        
        pdf.ln(20)
        pdf.set_font('DejaVu', '', 12)
        pdf.cell(0, 10, "توقيع الموظف: ________________", 0, 1, 'R')
        pdf.cell(0, 10, "توقيع المدير: ________________", 0, 1, 'R')
        pdf.cell(0, 10, f"التاريخ: {datetime.date.today().strftime('%Y-%m-%d')}", 0, 1, 'R')
        
        # Save to temp file
        temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        pdf_path = temp_file.name
        pdf.output(pdf_path)
        temp_file.close()
        
        return pdf_path

# ... (بقية الأكواد) ...

def main():
    app = QApplication(sys.argv)
    
    # Set application properties
    app.setApplicationName("نظام إدارة الموظفين - Perfection v3")
    app.setApplicationVersion("3.0")
    app.setOrganizationName("Perfection Systems")
    
    # Set application style
    app.setStyle('Fusion')
    
    # Set application palette
    palette = QPalette()
    palette.setColor(QPalette.Window, QColor(240, 240, 240))
    palette.setColor(QPalette.WindowText, QColor(44, 62, 80))
    palette.setColor(QPalette.Base, QColor(255, 255, 255))
    palette.setColor(QPalette.AlternateBase, QColor(245, 245, 245))
    palette.setColor(QPalette.ToolTipBase, QColor(255, 255, 255))
    palette.setColor(QPalette.ToolTipText, QColor(44, 62, 80))
    palette.setColor(QPalette.Text, QColor(44, 62, 80))
    palette.setColor(QPalette.Button, QColor(240, 240, 240))
    palette.setColor(QPalette.ButtonText, QColor(44, 62, 80))
    palette.setColor(QPalette.BrightText, QColor(255, 0, 0))
    palette.setColor(QPalette.Link, QColor(76, 175, 80))
    palette.setColor(QPalette.Highlight, QColor(76, 175, 80))
    palette.setColor(QPalette.HighlightedText, QColor(255, 255, 255))
    app.setPalette(palette)
    
    # Create and show main window
    window = EnhancedEmployeeSystem()
    
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()