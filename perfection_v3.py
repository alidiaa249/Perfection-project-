import json
import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import datetime
import pandas as pd
from tabulate import tabulate
import hashlib
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class EnhancedEmployeeSystem:
    def __init__(self, root):
        self.root = root
        self.root.title("perfection employee_system v3")
        self.root.geometry("1400x850")
        self.root.configure(bg='#f0f0f0')
        
        # Initialize data structures
        self.employees = {}
        self.other_employees = {}
        self.users = {"admin": self.hash_password("admin123")}
        self.current_user = None
        self.current_month = datetime.datetime.now().month
        self.current_year = datetime.datetime.now().year
        self.data_file = "employee_data.json"
        
        # Load saved data
        self.load_data()
        
        # Configure styles
        self.configure_styles()
        
        # Create login window
        self.create_login_window()
    
    def configure_styles(self):
        style = ttk.Style()
        style.theme_use('clam')
        
        # Configure colors
        style.configure('TFrame', background='#f0f0f0')
        style.configure('TLabel', background='#f0f0f0', font=('Arial', 12))
        style.configure('TButton', font=('Arial', 12), padding=5)
        style.configure('Accent.TButton', foreground='white', background='#4CAF50', font=('Arial', 12))
        style.configure('Edit.TButton', foreground='white', background='#2196F3', font=('Arial', 10))
        style.configure('Delete.TButton', foreground='white', background='#f44336', font=('Arial', 10))
        style.configure('Export.TButton', foreground='white', background='#FF9800', font=('Arial', 10))
        style.configure('TEntry', font=('Arial', 12))
        style.configure('TCombobox', font=('Arial', 12))
        style.configure('Treeview', font=('Arial', 11), rowheight=25)
        style.configure('Treeview.Heading', font=('Arial', 12, 'bold'))
        style.map('Treeview', background=[('selected', '#0078D7')])
    
    def hash_password(self, password):
        return hashlib.sha256(password.encode()).hexdigest()
    
    def save_data(self):
        """حفظ جميع البيانات في ملف JSON"""
        data = {
            'employees': self.employees,
            'other_employees': self.other_employees,
            'users': self.users
        }
        try:
            with open(self.data_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
        except Exception as e:
            messagebox.showerror("خطأ", f"فشل حفظ البيانات: {str(e)}")

    def load_data(self):
        """تحميل البيانات من ملف JSON"""
        if os.path.exists(self.data_file):
            try:
                with open(self.data_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.employees = data.get('employees', {})
                    self.other_employees = data.get('other_employees', {})
                    self.users = data.get('users', {"admin": self.hash_password("admin123")})
            except Exception as e:
                messagebox.showerror("خطأ", f"فشل تحميل البيانات: {str(e)}")
    
    def create_login_window(self):
        self.login_window = tk.Toplevel(self.root)
        self.login_window.title("تسجيل الدخول")
        self.login_window.geometry("400x250")
        self.login_window.configure(bg='#f0f0f0')
        
        tk.Label(self.login_window, text="perfection employee_system v3", font=('Arial', 16, 'bold'), bg='#f0f0f0').pack(pady=20)
        
        frame = ttk.Frame(self.login_window)
        frame.pack(pady=10)
        
        ttk.Label(frame, text="اسم المستخدم:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.username_entry = ttk.Entry(frame)
        self.username_entry.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(frame, text="كلمة المرور:").grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.password_entry = ttk.Entry(frame, show="*")
        self.password_entry.grid(row=1, column=1, padx=5, pady=5)
        
        btn_frame = ttk.Frame(self.login_window)
        btn_frame.pack(pady=20)
        
        ttk.Button(btn_frame, text="دخول", command=self.authenticate, 
                  style='Accent.TButton').pack(side='left', padx=10)
    
    def authenticate(self):
        username = self.username_entry.get()
        password = self.password_entry.get()
        
        if username in self.users and self.users[username] == self.hash_password(password):
            self.current_user = username
            self.login_window.destroy()
            self.create_main_interface()
        else:
            messagebox.showerror("خطأ", "اسم المستخدم أو كلمة المرور غير صحيحة")
    
    def create_main_interface(self):
        # Create notebook (tabs)
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True)
        
        # Create tabs
        self.create_employee_tab()
        self.create_attendance_tab()
        self.create_bonus_tab()
        self.create_advance_tab()
        self.create_reports_tab()
        
        # Load initial data
        self.update_employee_lists()
    
    def create_employee_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="إدارة الموظفين")
        
        # Regular Employees Frame
        reg_frame = ttk.LabelFrame(tab, text="موظفين بحصص", padding=10)
        reg_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        
        ttk.Label(reg_frame, text="الاسم:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.reg_name = ttk.Entry(reg_frame)
        self.reg_name.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(reg_frame, text="سعر الحصة الأساسي:").grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.reg_rate = ttk.Entry(reg_frame)
        self.reg_rate.grid(row=1, column=1, padx=5, pady=5)
        
        btn_frame = ttk.Frame(reg_frame)
        btn_frame.grid(row=2, column=0, columnspan=2, pady=10)
        
        ttk.Button(btn_frame, text="حفظ البيانات", command=self.save_regular_employee,
                  style='Accent.TButton').pack(side='left', padx=5)
        
        ttk.Button(btn_frame, text="حذف الموظف", command=self.delete_regular_employee,
                  style='Delete.TButton').pack(side='left', padx=5)
        
        # Other Employees Frame
        other_frame = ttk.LabelFrame(tab, text="موظفين براتب ثابت", padding=10)
        other_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")
        
        ttk.Label(other_frame, text="الاسم:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.other_name = ttk.Entry(other_frame)
        self.other_name.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(other_frame, text="الراتب الأساسي:").grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.other_salary = ttk.Entry(other_frame)
        self.other_salary.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Label(other_frame, text="الشهر:").grid(row=2, column=0, padx=5, pady=5, sticky='e')
        self.other_month = ttk.Combobox(other_frame, values=list(range(1, 13)), state="readonly")
        self.other_month.current(self.current_month - 1)
        self.other_month.grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Label(other_frame, text="السنة:").grid(row=3, column=0, padx=5, pady=5, sticky='e')
        self.other_year = ttk.Combobox(other_frame, values=list(range(2020, 2031)), state="readonly")
        self.other_year.current(datetime.datetime.now().year - 2020)
        self.other_year.grid(row=3, column=1, padx=5, pady=5)
        
        ttk.Label(other_frame, text="الراتب الشهري:").grid(row=4, column=0, padx=5, pady=5, sticky='e')
        self.other_monthly_salary = ttk.Entry(other_frame)
        self.other_monthly_salary.grid(row=4, column=1, padx=5, pady=5)
        
        btn_frame = ttk.Frame(other_frame)
        btn_frame.grid(row=5, column=0, columnspan=2, pady=10)
        
        ttk.Button(btn_frame, text="حفظ البيانات", command=self.save_other_employee,
                  style='Accent.TButton').pack(side='left', padx=5)
        
        ttk.Button(btn_frame, text="حذف الموظف", command=self.delete_other_employee,
                  style='Delete.TButton').pack(side='left', padx=5)
        
        ttk.Button(btn_frame, text="تحديث الراتب", command=self.update_salary,
                  style='Edit.TButton').pack(side='left', padx=5)
        
        # Employee List Frame
        list_frame = ttk.LabelFrame(tab, text="سجل الموظفين", padding=10)
        list_frame.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")
        
        columns = ("الاسم", "النوع", "سعر الحصة/الراتب", "حذف")
        self.employee_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=12)
        
        for col in columns:
            self.employee_tree.heading(col, text=col)
            self.employee_tree.column(col, width=120, anchor="center")
        
        self.employee_tree.pack(fill="both", expand=True)
        
        # Configure grid weights
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_columnconfigure(1, weight=1)
        tab.grid_rowconfigure(1, weight=1)
        
        # Bind delete buttons
        self.employee_tree.bind('<ButtonRelease-1>', self.handle_employee_tree_click)
    
    def handle_employee_tree_click(self, event):
        region = self.employee_tree.identify("region", event.x, event.y)
        if region != "cell":
            return
            
        column = self.employee_tree.identify_column(event.x)
        item = self.employee_tree.focus()
        col_index = int(column[1:]) - 1
        values = self.employee_tree.item(item, 'values')
        name = values[0]
        emp_type = values[1]
        
        if col_index == 3:  # Delete button column
            if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف الموظف {name}؟"):
                if emp_type == "بحصص":
                    if name in self.employees:
                        del self.employees[name]
                else:
                    if name in self.other_employees:
                        del self.other_employees[name]
                
                self.update_employee_lists()
                self.save_data()
                messagebox.showinfo("تم", f"تم حذف الموظف {name}")

    def create_attendance_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="تسجيل الحضور")
        
        # Attendance Entry Frame
        entry_frame = ttk.LabelFrame(tab, text="تسجيل يومي", padding=10)
        entry_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        
        ttk.Label(entry_frame, text="التاريخ (YYYY-MM-DD):").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.att_date = ttk.Entry(entry_frame)
        self.att_date.insert(0, datetime.date.today().strftime("%Y-%m-%d"))
        self.att_date.grid(row=0, column=1, padx=5, pady=5)
        
        # Add day name label
        ttk.Label(entry_frame, text="اليوم:").grid(row=0, column=2, padx=5, pady=5, sticky='e')
        self.day_name = ttk.Label(entry_frame, text="")
        self.day_name.grid(row=0, column=3, padx=5, pady=5, sticky='w')
        
        # Add reset date button
        ttk.Button(entry_frame, text="إعادة تعيين", command=self.reset_attendance_date,
                  style='Edit.TButton').grid(row=0, column=4, padx=5, pady=5)
        
        # Bind date change to update day name
        self.att_date.bind('<FocusOut>', self.update_day_name)
        
        btn_frame = ttk.Frame(entry_frame)
        btn_frame.grid(row=1, column=0, columnspan=5, pady=10)
        
        ttk.Button(btn_frame, text="تسجيل حضور لليوم", command=self.open_daily_attendance_window,
                  style='Accent.TButton').pack(side='left', padx=5)
        
        ttk.Button(btn_frame, text="تعديل حضور", command=self.edit_attendance,
                  style='Edit.TButton').pack(side='left', padx=5)
        
        # Month filter for attendance
        filter_frame = ttk.LabelFrame(tab, text="تصفية حسب الشهر", padding=10)
        filter_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")
        
        ttk.Label(filter_frame, text="الشهر:").grid(row=0, column=0, padx=5, pady=5)
        self.filter_month = ttk.Combobox(filter_frame, values=list(range(1, 13)), state="readonly")
        self.filter_month.current(self.current_month - 1)
        self.filter_month.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(filter_frame, text="السنة:").grid(row=1, column=0, padx=5, pady=5)
        self.filter_year = ttk.Combobox(filter_frame, values=list(range(2020, 2031)), state="readonly")
        self.filter_year.current(datetime.datetime.now().year - 2020)
        self.filter_year.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Button(filter_frame, text="تصفية", command=self.filter_attendance_by_month,
                  style='Accent.TButton').grid(row=2, column=0, columnspan=2, pady=5)
        
        ttk.Button(filter_frame, text="إعادة تعيين", command=self.reset_attendance_filter,
                  style='Accent.TButton').grid(row=3, column=0, columnspan=2, pady=5)
        
        # Attendance List Frame
        list_frame = ttk.LabelFrame(tab, text="سجل الحضور", padding=10)
        list_frame.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")
        
        columns = ("الموظف", "التاريخ", "اليوم", "الحصص", "البونص اليومي", "تعديل", "حذف")
        self.attendance_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=15)
        
        for col in columns:
            self.attendance_tree.heading(col, text=col)
            self.attendance_tree.column(col, width=100, anchor="center")
        
        self.attendance_tree.pack(fill="both", expand=True)
        
        # Configure grid weights
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_columnconfigure(1, weight=1)
        tab.grid_rowconfigure(1, weight=1)
        
        # Bind delete buttons
        self.attendance_tree.bind('<ButtonRelease-1>', self.handle_attendance_tree_click)
    
    def reset_attendance_date(self):
        self.att_date.delete(0, tk.END)
        self.att_date.insert(0, datetime.date.today().strftime("%Y-%m-%d"))
        self.update_day_name()
    
    def edit_attendance(self):
        selected_item = self.attendance_tree.focus()
        if not selected_item:
            messagebox.showerror("خطأ", "الرجاء تحديد سجل حضور لتعديله")
            return
        
        values = self.attendance_tree.item(selected_item, 'values')
        date = values[1]
        
        # فتح نافذة تسجيل الحضور مع البيانات الحالية
        self.att_date.delete(0, tk.END)
        self.att_date.insert(0, date)
        self.open_daily_attendance_window(edit_mode=True)
    
    def open_daily_attendance_window(self, edit_mode=False):
        date = self.att_date.get()
        
        try:
            datetime.datetime.strptime(date, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("خطأ", "صيغة التاريخ غير صحيحة (يجب أن تكون YYYY-MM-DD)")
            return
        
        if not self.employees:
            messagebox.showerror("خطأ", "لا يوجد موظفين بحصص مسجلين")
            return
        
        # Create new window for daily attendance
        daily_window = tk.Toplevel(self.root)
        daily_window.title(f"تسجيل حضور ليوم {date}")
        daily_window.geometry("800x600")
        
        # Create a frame for the attendance entries
        frame = ttk.Frame(daily_window)
        frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Create a canvas and scrollbar
        canvas = tk.Canvas(frame)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Add headers
        headers_frame = ttk.Frame(scrollable_frame)
        headers_frame.pack(fill='x', pady=5)
        
        ttk.Label(headers_frame, text="الموظف", font=('Arial', 12, 'bold')).pack(side='left', padx=10, expand=True)
        ttk.Label(headers_frame, text="عدد الحصص", font=('Arial', 12, 'bold')).pack(side='left', padx=10, expand=True)
        ttk.Label(headers_frame, text="البونص اليومي", font=('Arial', 12, 'bold')).pack(side='left', padx=10, expand=True)
        
        # Create entry widgets for each employee
        self.daily_att_entries = {}
        
        for name in self.employees:
            emp_frame = ttk.Frame(scrollable_frame)
            emp_frame.pack(fill='x', pady=5)
            
            ttk.Label(emp_frame, text=name).pack(side='left', padx=10, expand=True)
            
            sessions_var = tk.StringVar(value="0")
            bonus_var = tk.StringVar(value="0")
            
            # إذا كان في وضع التعديل، نملأ البيانات الحالية
            if edit_mode and 'attendance' in self.employees[name] and date in self.employees[name]['attendance']:
                sessions_var.set(str(self.employees[name]['attendance'][date].get('sessions', 0)))
                bonus_var.set(str(self.employees[name]['attendance'][date].get('daily_bonus', 0)))
            
            sessions_entry = ttk.Entry(emp_frame, textvariable=sessions_var)
            sessions_entry.pack(side='left', padx=10, expand=True)
            
            bonus_entry = ttk.Entry(emp_frame, textvariable=bonus_var)
            bonus_entry.pack(side='left', padx=10, expand=True)
            
            self.daily_att_entries[name] = {
                'sessions': sessions_var,
                'bonus': bonus_var
            }
        
        # Add save button
        btn_frame = ttk.Frame(daily_window)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="حفظ الحضور", command=lambda: self.save_daily_attendance(date, daily_window, edit_mode),
                  style='Accent.TButton').pack(side='left', padx=10)
        
        ttk.Button(btn_frame, text="إلغاء", command=daily_window.destroy,
                  style='Delete.TButton').pack(side='left', padx=10)
    
    def save_daily_attendance(self, date, window, edit_mode=False):
        try:
            date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("خطأ", "صيغة التاريخ غير صحيحة")
            return
        
        for name, entries in self.daily_att_entries.items():
            sessions = entries['sessions'].get()
            bonus = entries['bonus'].get()
            
            try:
                sessions = int(sessions) if sessions else 0
                bonus = float(bonus) if bonus else 0
            except ValueError:
                messagebox.showerror("خطأ", f"قيم غير صحيحة للموظف {name}")
                return
            
            # إذا كان عدد الحصص صفر، لا نضيف السجل
            if sessions == 0:
                # حذف السجل إذا كان موجودًا
                if name in self.employees and 'attendance' in self.employees[name] and date in self.employees[name]['attendance']:
                    del self.employees[name]['attendance'][date]
                continue
            
            if name not in self.employees:
                continue
            
            if 'attendance' not in self.employees[name]:
                self.employees[name]['attendance'] = {}
            
            self.employees[name]['attendance'][date] = {
                'sessions': sessions,
                'daily_bonus': bonus
            }
        
        self.save_data()
        self.update_employee_lists()
        message = "تم تعديل سجل الحضور" if edit_mode else "تم تسجيل حضور لليوم"
        messagebox.showinfo("تم", f"{message} {date}")
        window.destroy()
    
    def update_day_name(self, event=None):
        try:
            date_str = self.att_date.get()
            date_obj = datetime.datetime.strptime(date_str, "%Y-%m-%d")
            # Arabic day names
            days = ["الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت", "الأحد"]
            day_index = date_obj.weekday()
            self.day_name.config(text=days[day_index])
        except ValueError:
            self.day_name.config(text="")
    
    def filter_attendance_by_month(self):
        month = int(self.filter_month.get())
        year = int(self.filter_year.get())
        
        # تحديد بداية ونهاية الشهر
        month_start = datetime.datetime(year, month, 1)
        month_end = datetime.datetime(year, month, 1) + datetime.timedelta(days=32)
        month_end = month_end.replace(day=1) - datetime.timedelta(days=1)
        
        self.attendance_tree.delete(*self.attendance_tree.get_children())
        
        # ترتيب الموظفين حسب الاسم
        sorted_employees = sorted(self.employees.items(), key=lambda x: x[0])
        
        for name, data in sorted_employees:
            # ترتيب سجلات الحضور حسب التاريخ
            sorted_attendance = sorted(data.get('attendance', {}).items(), key=lambda x: x[0])
            
            for date, record in sorted_attendance:
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                
                # Apply month filter
                if date_obj < month_start or date_obj > month_end:
                    continue
                
                day_name = ["الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت", "الأحد"][date_obj.weekday()]
                
                self.attendance_tree.insert("", "end", values=(
                    name,
                    date,
                    day_name,
                    record.get('sessions', 0),
                    f"{record.get('daily_bonus', 0):.2f}",
                    "تعديل",
                    "حذف"
                ), tags=('editable',))
    
    def reset_attendance_filter(self):
        self.filter_month.current(self.current_month - 1)
        self.filter_year.current(datetime.datetime.now().year - 2020)
        self.update_employee_lists()
    
    def handle_attendance_tree_click(self, event):
        region = self.attendance_tree.identify("region", event.x, event.y)
        if region != "cell":
            return
            
        column = self.attendance_tree.identify_column(event.x)
        item = self.attendance_tree.focus()
        col_index = int(column[1:]) - 1
        values = self.attendance_tree.item(item, 'values')
        name = values[0]
        date = values[1]
        
        if col_index == 6:  # Delete button column
            if name in self.employees and date in self.employees[name].get('attendance', {}):
                if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف تسجيل حضور {name} بتاريخ {date}؟"):
                    del self.employees[name]['attendance'][date]
                    self.update_employee_lists()
                    self.save_data()
                    messagebox.showinfo("تم", "تم حذف تسجيل الحضور")

    def create_bonus_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="المكافآت والخصومات")
        
        # Monthly Bonus Frame
        bonus_frame = ttk.LabelFrame(tab, text="تسجيل مكافآت الأداء", padding=10)
        bonus_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        
        ttk.Label(bonus_frame, text="الموظف:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.bonus_employee = ttk.Combobox(bonus_frame, values=[], state="readonly")
        self.bonus_employee.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(bonus_frame, text="التاريخ (YYYY-MM-DD):").grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.bonus_date = ttk.Entry(bonus_frame)
        self.bonus_date.insert(0, datetime.date.today().strftime("%Y-%m-%d"))
        self.bonus_date.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Label(bonus_frame, text="عدد حصص الأداء:").grid(row=2, column=0, padx=5, pady=5, sticky='e')
        self.performance_bonus_sessions = ttk.Entry(bonus_frame)
        self.performance_bonus_sessions.insert(0, "0")
        self.performance_bonus_sessions.grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Label(bonus_frame, text="سعر الحصة:").grid(row=3, column=0, padx=5, pady=5, sticky='e')
        self.current_month_rate = ttk.Entry(bonus_frame)
        self.current_month_rate.grid(row=3, column=1, padx=5, pady=5)
        
        ttk.Label(bonus_frame, text="البونص الشهري:").grid(row=4, column=0, padx=5, pady=5, sticky='e')
        self.monthly_bonus = ttk.Entry(bonus_frame)
        self.monthly_bonus.insert(0, "0")
        self.monthly_bonus.grid(row=4, column=1, padx=5, pady=5)
        
        btn_frame = ttk.Frame(bonus_frame)
        btn_frame.grid(row=5, column=0, columnspan=2, pady=10)
        
        ttk.Button(btn_frame, text="حفظ المكافآت", command=self.save_bonus,
                  style='Accent.TButton').pack(side='left', padx=5)
        
        ttk.Button(btn_frame, text="حذف المكافآت", command=self.delete_bonus,
                  style='Delete.TButton').pack(side='left', padx=5)
        
        ttk.Button(btn_frame, text="عرض سجل المكافآت", command=self.show_bonus_history,
                  style='Accent.TButton').pack(side='left', padx=5)
        
        ttk.Button(btn_frame, text="تسجيل جماعي", command=self.open_collective_bonus_window,
                  style='Accent.TButton').pack(side='left', padx=5)
        
        # Deductions Frame
        ded_frame = ttk.LabelFrame(tab, text="تسجيل الخصومات", padding=10)
        ded_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")
        
        ttk.Label(ded_frame, text="الموظف:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.ded_employee = ttk.Combobox(ded_frame, values=[], state="readonly")
        self.ded_employee.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(ded_frame, text="التاريخ (YYYY-MM-DD):").grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.ded_date = ttk.Entry(ded_frame)
        self.ded_date.insert(0, datetime.date.today().strftime("%Y-%m-%d"))
        self.ded_date.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Label(ded_frame, text="قيمة الخصم:").grid(row=2, column=0, padx=5, pady=5, sticky='e')
        self.ded_amount = ttk.Entry(ded_frame)
        self.ded_amount.insert(0, "0")
        self.ded_amount.grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Label(ded_frame, text="السبب:").grid(row=3, column=0, padx=5, pady=5, sticky='e')
        self.ded_reason = ttk.Entry(ded_frame)
        self.ded_reason.grid(row=3, column=1, padx=5, pady=5)
        
        btn_frame = ttk.Frame(ded_frame)
        btn_frame.grid(row=4, column=0, columnspan=2, pady=10)
        
        ttk.Button(btn_frame, text="حفظ الخصم", command=self.save_deduction,
                  style='Accent.TButton').pack(side='left', padx=5)
        
        ttk.Button(btn_frame, text="حذف الخصم", command=self.delete_deduction,
                  style='Delete.TButton').pack(side='left', padx=5)
        
        ttk.Button(btn_frame, text="عرض سجل الخصومات", command=self.show_deduction_history,
                  style='Accent.TButton').pack(side='left', padx=5)
        
        # Bonus List Frame
        list_frame = ttk.LabelFrame(tab, text="سجل المكافآت", padding=10)
        list_frame.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")
        
        columns = ("الموظف", "التاريخ", "نوع المكافأة", "التفاصيل", "القيمة", "تعديل", "حذف")
        self.bonus_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=15)
        
        for col in columns:
            self.bonus_tree.heading(col, text=col)
            self.bonus_tree.column(col, width=100, anchor="center")
        
        self.bonus_tree.pack(fill="both", expand=True)
        
        # Configure grid weights
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_columnconfigure(1, weight=1)
        tab.grid_rowconfigure(1, weight=1)
        
        # Bind delete buttons
        self.bonus_tree.bind('<ButtonRelease-1>', self.handle_bonus_tree_click)
    
    def open_collective_bonus_window(self):
        date = self.bonus_date.get()
        
        try:
            datetime.datetime.strptime(date, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("خطأ", "صيغة التاريخ غير صحيحة (يجب أن تكون YYYY-MM-DD)")
            return
        
        if not self.employees:
            messagebox.showerror("خطأ", "لا يوجد موظفين بحصص مسجلين")
            return
        
        # Create new window for collective bonus
        collective_window = tk.Toplevel(self.root)
        collective_window.title(f"تسجيل مكافآت جماعية ليوم {date}")
        collective_window.geometry("800x600")
        
        # Create a frame for the bonus entries
        frame = ttk.Frame(collective_window)
        frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Create a canvas and scrollbar
        canvas = tk.Canvas(frame)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Add headers
        headers_frame = ttk.Frame(scrollable_frame)
        headers_frame.pack(fill='x', pady=5)
        
        ttk.Label(headers_frame, text="الموظف", font=('Arial', 12, 'bold')).pack(side='left', padx=10, expand=True)
        ttk.Label(headers_frame, text="حصص الأداء", font=('Arial', 12, 'bold')).pack(side='left', padx=10, expand=True)
        ttk.Label(headers_frame, text="سعر الحصة", font=('Arial', 12, 'bold')).pack(side='left', padx=10, expand=True)
        ttk.Label(headers_frame, text="البونص الشهري", font=('Arial', 12, 'bold')).pack(side='left', padx=10, expand=True)
        
        # Create entry widgets for each employee
        self.collective_bonus_entries = {}
        
        for name in self.employees:
            emp_frame = ttk.Frame(scrollable_frame)
            emp_frame.pack(fill='x', pady=5)
            
            ttk.Label(emp_frame, text=name).pack(side='left', padx=10, expand=True)
            
            sessions_var = tk.StringVar(value="0")
            rate_var = tk.StringVar(value=str(self.employees[name].get('current_rate', 0)))
            bonus_var = tk.StringVar(value="0")
            
            sessions_entry = ttk.Entry(emp_frame, textvariable=sessions_var)
            sessions_entry.pack(side='left', padx=10, expand=True)
            
            rate_entry = ttk.Entry(emp_frame, textvariable=rate_var)
            rate_entry.pack(side='left', padx=10, expand=True)
            
            bonus_entry = ttk.Entry(emp_frame, textvariable=bonus_var)
            bonus_entry.pack(side='left', padx=10, expand=True)
            
            self.collective_bonus_entries[name] = {
                'sessions': sessions_var,
                'rate': rate_var,
                'bonus': bonus_var
            }
        
        # Add save button
        btn_frame = ttk.Frame(collective_window)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="حفظ المكافآت", command=lambda: self.save_collective_bonus(date, collective_window),
                  style='Accent.TButton').pack(side='left', padx=10)
        
        ttk.Button(btn_frame, text="إلغاء", command=collective_window.destroy,
                  style='Delete.TButton').pack(side='left', padx=10)
    
    def save_collective_bonus(self, date, window):
        try:
            date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("خطأ", "صيغة التاريخ غير صحيحة")
            return
        
        for name, entries in self.collective_bonus_entries.items():
            sessions = entries['sessions'].get()
            rate = entries['rate'].get()
            bonus = entries['bonus'].get()
            
            try:
                sessions = int(sessions) if sessions else 0
                rate = float(rate) if rate else 0
                bonus = float(bonus) if bonus else 0
            except ValueError:
                messagebox.showerror("خطأ", f"قيم غير صحيحة للموظف {name}")
                return
            
            if name not in self.employees:
                continue
            
            # حساب بونص الأداء (عدد الحصص × سعر الحصة)
            performance_bonus_amount = sessions * rate
            
            # حفظ بونص الأداء إذا كان هناك حصص
            if sessions > 0:
                if 'performance_bonus' not in self.employees[name]:
                    self.employees[name]['performance_bonus'] = {}
                
                self.employees[name]['performance_bonus'][date] = {
                    'sessions': sessions,
                    'amount': performance_bonus_amount,
                    'rate': rate
                }
            
            # حفظ البونص الشهري إذا كان هناك قيمة
            if bonus > 0:
                if 'monthly_bonuses' not in self.employees[name]:
                    self.employees[name]['monthly_bonuses'] = {}
                
                # نستخدم التاريخ كمعرف للبونص الشهري
                self.employees[name]['monthly_bonuses'][date] = bonus
            
            # حفظ سعر الحصة إذا كان مختلفًا عن السعر الحالي
            if rate > 0 and rate != self.employees[name].get('current_rate', 0):
                if 'monthly_rates' not in self.employees[name]:
                    self.employees[name]['monthly_rates'] = {}
                
                # نستخدم التاريخ كمعرف لسعر الحصة
                self.employees[name]['monthly_rates'][date] = rate
                self.employees[name]['current_rate'] = rate
        
        self.save_data()
        self.update_employee_lists()
        messagebox.showinfo("تم", f"تم تسجيل المكافآت الجماعية لليوم {date}")
        window.destroy()
    
    def handle_bonus_tree_click(self, event):
        region = self.bonus_tree.identify("region", event.x, event.y)
        if region != "cell":
            return
            
        column = self.bonus_tree.identify_column(event.x)
        item = self.bonus_tree.focus()
        col_index = int(column[1:]) - 1
        values = self.bonus_tree.item(item, 'values')
        name = values[0]
        date = values[1]
        bonus_type = values[2]
        
        if col_index == 6:  # Delete button column
            if name in self.employees:
                if bonus_type == "بونص الأداء" and date in self.employees[name].get('performance_bonus', {}):
                    if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف مكافأة الأداء لـ {name} بتاريخ {date}؟"):
                        del self.employees[name]['performance_bonus'][date]
                        self.update_employee_lists()
                        self.save_data()
                        messagebox.showinfo("تم", "تم حذف المكافأة")
                elif bonus_type == "البونص الشهري" and date in self.employees[name].get('monthly_bonuses', {}):
                    if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف البونص الشهري لـ {name} بتاريخ {date}؟"):
                        del self.employees[name]['monthly_bonuses'][date]
                        self.update_employee_lists()
                        self.save_data()
                        messagebox.showinfo("تم", "تم حذف المكافأة")
    
    def save_bonus(self):
        name = self.bonus_employee.get()
        date = self.bonus_date.get()
        performance_bonus_sessions = self.performance_bonus_sessions.get()
        monthly_bonus = self.monthly_bonus.get()
        current_month_rate = self.current_month_rate.get()
        
        if not name or not date:
            messagebox.showerror("خطأ", "الموظف والتاريخ مطلوبان")
            return
        
        try:
            date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
            performance_bonus_sessions = int(performance_bonus_sessions) if performance_bonus_sessions else 0
            monthly_bonus = float(monthly_bonus) if monthly_bonus else 0
            current_month_rate = float(current_month_rate) if current_month_rate else 0
        except ValueError:
            messagebox.showerror("خطأ", "الرجاء إدخال أرقام صحيحة")
            return
        
        if name not in self.employees:
            messagebox.showerror("خطأ", "الموظف غير موجود")
            return
        
        # حساب بونص الأداء (عدد الحصص × سعر الحصة)
        performance_bonus_amount = performance_bonus_sessions * current_month_rate
        
        # حفظ البيانات
        if 'performance_bonus' not in self.employees[name]:
            self.employees[name]['performance_bonus'] = {}
        
        if 'monthly_rates' not in self.employees[name]:
            self.employees[name]['monthly_rates'] = {}
        
        if 'monthly_bonuses' not in self.employees[name]:
            self.employees[name]['monthly_bonuses'] = {}
        
        # حفظ بونص الأداء
        if performance_bonus_sessions > 0:
            self.employees[name]['performance_bonus'][date] = {
                'sessions': performance_bonus_sessions,
                'amount': performance_bonus_amount,
                'rate': current_month_rate
            }
        
        # حفظ البونص الشهري
        if monthly_bonus > 0:
            self.employees[name]['monthly_bonuses'][date] = monthly_bonus
        
        # حفظ سعر الحصة إذا كان مختلفًا عن السعر الحالي
        if current_month_rate > 0 and current_month_rate != self.employees[name].get('current_rate', 0):
            self.employees[name]['monthly_rates'][date] = current_month_rate
            self.employees[name]['current_rate'] = current_month_rate
        
        self.save_data()
        messagebox.showinfo("تم", f"تم حفظ المكافآت لـ {name}\nبونص الأداء: {performance_bonus_amount:.2f} (عدد الحصص: {performance_bonus_sessions})\nالبونص الشهري: {monthly_bonus:.2f}")
        self.clear_bonus_fields()
        self.update_employee_lists()
    
    def delete_bonus(self):
        name = self.bonus_employee.get()
        date = self.bonus_date.get()
        
        if not name or not date:
            messagebox.showerror("خطأ", "الرجاء تحديد الموظف والتاريخ")
            return
        
        if name not in self.employees:
            messagebox.showerror("خطأ", "الموظف غير موجود")
            return
        
        if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف مكافآت {name} بتاريخ {date}؟"):
            if 'performance_bonus' in self.employees[name] and date in self.employees[name]['performance_bonus']:
                del self.employees[name]['performance_bonus'][date]
            
            if 'monthly_bonuses' in self.employees[name] and date in self.employees[name]['monthly_bonuses']:
                del self.employees[name]['monthly_bonuses'][date]
            
            self.save_data()
            messagebox.showinfo("تم", "تم حذف المكافآت")
            self.clear_bonus_fields()
            self.update_employee_lists()
    
    def show_bonus_history(self):
        name = self.bonus_employee.get()
        if not name:
            messagebox.showerror("خطأ", "الرجاء اختيار موظف")
            return
        
        if name not in self.employees:
            messagebox.showerror("خطأ", "هذا الموظف ليس من الموظفين بحصص")
            return
        
        bonus_history = self.employees[name].get('performance_bonus', {})
        monthly_bonuses = self.employees[name].get('monthly_bonuses', {})
        
        if not bonus_history and not monthly_bonuses:
            messagebox.showinfo("السجل", "لا يوجد سجل مكافآت لهذا الموظف")
            return
        
        # Update bonus tree
        self.bonus_tree.delete(*self.bonus_tree.get_children())
        
        # Add performance bonuses
        for date in sorted(bonus_history.keys()):
            sessions = bonus_history[date].get('sessions', 0)
            rate = bonus_history[date].get('rate', 0)
            amount = bonus_history[date].get('amount', 0)
            
            self.bonus_tree.insert("", "end", values=(
                name,
                date,
                "بونص الأداء",
                f"{sessions} حصة × {rate:.2f}",
                f"{amount:.2f}",
                "تعديل",
                "حذف"
            ))
        
        # Add monthly bonuses
        for date in sorted(monthly_bonuses.keys()):
            bonus = monthly_bonuses[date]
            if bonus > 0:
                self.bonus_tree.insert("", "end", values=(
                    name,
                    date,
                    "البونص الشهري",
                    "-",
                    f"{bonus:.2f}",
                    "تعديل",
                    "حذف"
                ))
    
    def save_deduction(self):
        name = self.ded_employee.get()
        date = self.ded_date.get()
        amount = self.ded_amount.get()
        reason = self.ded_reason.get()
        
        if not name or not date or not amount:
            messagebox.showerror("خطأ", "الموظف والتاريخ وقيمة الخصم مطلوبة")
            return
        
        try:
            date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
            amount = float(amount)
        except ValueError:
            messagebox.showerror("خطأ", "الرجاء إدخال أرقام صحيحة")
            return
        
        if name in self.employees:
            if 'deductions' not in self.employees[name]:
                self.employees[name]['deductions'] = {}
            
            self.employees[name]['deductions'][date] = {
                'amount': amount,
                'reason': reason
            }
        elif name in self.other_employees:
            if 'deductions' not in self.other_employees[name]:
                self.other_employees[name]['deductions'] = {}
            
            self.other_employees[name]['deductions'][date] = {
                'amount': amount,
                'reason': reason
            }
        else:
            messagebox.showerror("خطأ", "الموظف غير موجود")
            return
        
        self.save_data()
        messagebox.showinfo("تم", f"تم حفظ الخصم لـ {name} بتاريخ {date}")
        self.clear_deduction_fields()
    
    def delete_deduction(self):
        name = self.ded_employee.get()
        date = self.ded_date.get()
        
        if not name or not date:
            messagebox.showerror("خطأ", "الرجاء تحديد الموظف والتاريخ")
            return
        
        if name in self.employees and date in self.employees[name].get('deductions', {}):
            if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف خصم {name} بتاريخ {date}؟"):
                del self.employees[name]['deductions'][date]
                self.save_data()
                messagebox.showinfo("تم", "تم حذف الخصم")
                self.clear_deduction_fields()
        elif name in self.other_employees and date in self.other_employees[name].get('deductions', {}):
            if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف خصم {name} بتاريخ {date}؟"):
                del self.other_employees[name]['deductions'][date]
                self.save_data()
                messagebox.showinfo("تم", "تم حذف الخصم")
                self.clear_deduction_fields()
        else:
            messagebox.showerror("خطأ", "الخصم غير موجود")
    
    def show_deduction_history(self):
        name = self.ded_employee.get()
        if not name:
            messagebox.showerror("خطأ", "الرجاء اختيار موظف")
            return
        
        if name in self.employees:
            deductions = self.employees[name].get('deductions', {})
        elif name in self.other_employees:
            deductions = self.other_employees[name].get('deductions', {})
        else:
            messagebox.showerror("خطأ", "الموظف غير موجود")
            return
        
        if not deductions:
            messagebox.showinfo("السجل", "لا يوجد سجل خصومات لهذا الموظف")
            return
        
        history_text = f"سجل خصومات الموظف {name}:\n"
        history_text += "-"*80 + "\n"
        history_text += "التاريخ | المبلغ | السبب\n"
        history_text += "-"*80 + "\n"
        
        for date in sorted(deductions.keys()):
            amount = deductions[date].get('amount', 0)
            reason = deductions[date].get('reason', '')
            history_text += f"{date} | {amount:.2f} | {reason}\n"
        
        # عرض النتائج في نافذة جديدة
        history_window = tk.Toplevel(self.root)
        history_window.title(f"سجل خصومات الموظف {name}")
        
        text_widget = tk.Text(history_window, wrap="word", width=80, height=15)
        text_widget.pack(padx=10, pady=10)
        text_widget.insert("end", history_text)
        text_widget.config(state="disabled")
        
        # زر لحذف السجل
        def delete_selected():
            selected_date = date_var.get()
            if not selected_date:
                return
            
            if messagebox.askyesno("تأكيد", f"هل تريد حذف خصم بتاريخ {selected_date}؟"):
                if name in self.employees and selected_date in self.employees[name].get('deductions', {}):
                    del self.employees[name]['deductions'][selected_date]
                elif name in self.other_employees and selected_date in self.other_employees[name].get('deductions', {}):
                    del self.other_employees[name]['deductions'][selected_date]
                self.save_data()
                messagebox.showinfo("تم", "تم حذف الخصم")
                history_window.destroy()
        
        # قائمة اختيار التاريخ
        date_var = tk.StringVar()
        dates_list = sorted(deductions.keys())
        date_menu = ttk.Combobox(history_window, textvariable=date_var, values=dates_list, state="readonly")
        date_menu.pack(pady=5)
        
        ttk.Button(history_window, text="حذف المحدد", command=delete_selected,
                  style='Delete.TButton').pack(pady=5)
    
    def clear_bonus_fields(self):
        self.performance_bonus_sessions.delete(0, tk.END)
        self.performance_bonus_sessions.insert(0, "0")
        self.monthly_bonus.delete(0, tk.END)
        self.monthly_bonus.insert(0, "0")
        self.current_month_rate.delete(0, tk.END)
    
    def clear_deduction_fields(self):
        self.ded_amount.delete(0, tk.END)
        self.ded_amount.insert(0, "0")
        self.ded_reason.delete(0, tk.END)

    def create_advance_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="السلف")
        
        # Advance Entry Frame
        entry_frame = ttk.LabelFrame(tab, text="تسجيل سلفة", padding=10)
        entry_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        
        ttk.Label(entry_frame, text="الموظف:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.adv_employee = ttk.Combobox(entry_frame, values=[], state="readonly")
        self.adv_employee.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(entry_frame, text="التاريخ (YYYY-MM-DD):").grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.adv_date = ttk.Entry(entry_frame)
        self.adv_date.insert(0, datetime.date.today().strftime("%Y-%m-%d"))
        self.adv_date.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Label(entry_frame, text="المبلغ:").grid(row=2, column=0, padx=5, pady=5, sticky='e')
        self.adv_amount = ttk.Entry(entry_frame)
        self.adv_amount.grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Label(entry_frame, text="الشهر المستحق:").grid(row=3, column=0, padx=5, pady=5, sticky='e')
        self.adv_due_month = ttk.Combobox(entry_frame, values=list(range(1, 13)), state="readonly")
        self.adv_due_month.current(self.current_month - 1)
        self.adv_due_month.grid(row=3, column=1, padx=5, pady=5)
        
        ttk.Label(entry_frame, text="السنة المستحقة:").grid(row=4, column=0, padx=5, pady=5, sticky='e')
        self.adv_due_year = ttk.Combobox(entry_frame, values=list(range(2020, 2031)), state="readonly")
        self.adv_due_year.current(datetime.datetime.now().year - 2020)
        self.adv_due_year.grid(row=4, column=1, padx=5, pady=5)
        
        btn_frame = ttk.Frame(entry_frame)
        btn_frame.grid(row=5, column=0, columnspan=2, pady=10)
        
        ttk.Button(btn_frame, text="حفظ السلفة", command=self.save_advance,
                  style='Accent.TButton').pack(side='left', padx=5)
        
        ttk.Button(btn_frame, text="حذف السلفة", command=self.delete_advance,
                  style='Delete.TButton').pack(side='left', padx=5)
        
        # Advance List Frame
        list_frame = ttk.LabelFrame(tab, text="سجل السلف", padding=10)
        list_frame.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")
        
        columns = ("الموظف", "التاريخ", "المبلغ", "الشهر المستحق", "السنة المستحقة", "حذف")
        self.advance_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=15)
        
        for col in columns:
            self.advance_tree.heading(col, text=col)
            self.advance_tree.column(col, width=100, anchor="center")
        
        self.advance_tree.pack(fill="both", expand=True)
        
        # Configure grid weights
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_rowconfigure(1, weight=1)
        
        # Bind delete buttons
        self.advance_tree.bind('<ButtonRelease-1>', self.handle_advance_tree_click)
    
    def handle_advance_tree_click(self, event):
        region = self.advance_tree.identify("region", event.x, event.y)
        if region != "cell":
            return
            
        column = self.advance_tree.identify_column(event.x)
        item = self.advance_tree.focus()
        col_index = int(column[1:]) - 1
        values = self.advance_tree.item(item, 'values')
        name = values[0]
        date = values[1]
        
        if col_index == 5:  # Delete button column
            if name in self.employees and date in self.employees[name].get('advances', {}):
                if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف سلفة {name} بتاريخ {date}؟"):
                    del self.employees[name]['advances'][date]
                    if 'advance_due_dates' in self.employees[name]:
                        del self.employees[name]['advance_due_dates'][date]
                    self.update_employee_lists()
                    self.save_data()
                    messagebox.showinfo("تم", "تم حذف السلفة")
            elif name in self.other_employees and date in self.other_employees[name].get('advances', {}):
                if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف سلفة {name} بتاريخ {date}؟"):
                    del self.other_employees[name]['advances'][date]
                    if 'advance_due_dates' in self.other_employees[name]:
                        del self.other_employees[name]['advance_due_dates'][date]
                    self.update_employee_lists()
                    self.save_data()
                    messagebox.showinfo("تم", "تم حذف السلفة")
    
    def save_advance(self):
        name = self.adv_employee.get()
        date = self.adv_date.get()
        amount = self.adv_amount.get()
        due_month = self.adv_due_month.get()
        due_year = self.adv_due_year.get()
        
        if not name or not date or not amount or not due_month or not due_year:
            messagebox.showerror("خطأ", "جميع الحقول مطلوبة")
            return
        
        try:
            datetime.datetime.strptime(date, "%Y-%m-%d")
            amount = float(amount)
            due_month = int(due_month)
            due_year = int(due_year)
        except ValueError:
            messagebox.showerror("خطأ", "صيغة التاريخ أو القيم غير صحيحة")
            return
        
        if name in self.employees:
            if 'advances' not in self.employees[name]:
                self.employees[name]['advances'] = {}
            if 'advance_due_dates' not in self.employees[name]:
                self.employees[name]['advance_due_dates'] = {}
            
            self.employees[name]['advances'][date] = amount
            self.employees[name]['advance_due_dates'][date] = {
                'month': due_month,
                'year': due_year
            }
        elif name in self.other_employees:
            if 'advances' not in self.other_employees[name]:
                self.other_employees[name]['advances'] = {}
            if 'advance_due_dates' not in self.other_employees[name]:
                self.other_employees[name]['advance_due_dates'] = {}
            
            self.other_employees[name]['advances'][date] = amount
            self.other_employees[name]['advance_due_dates'][date] = {
                'month': due_month,
                'year': due_year
            }
        else:
            messagebox.showerror("خطأ", "الموظف غير موجود")
            return
        
        self.update_employee_lists()
        self.save_data()
        messagebox.showinfo("تم", f"تم تسجيل سلفة لـ {name} بتاريخ {date}")
        self.clear_advance_fields()
    
    def delete_advance(self):
        name = self.adv_employee.get()
        date = self.adv_date.get()
        
        if not name or not date:
            messagebox.showerror("خطأ", "الرجاء تحديد الموظف والتاريخ")
            return
        
        if name in self.employees and date in self.employees[name].get('advances', {}):
            if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف سلفة {name} بتاريخ {date}؟"):
                del self.employees[name]['advances'][date]
                if 'advance_due_dates' in self.employees[name]:
                    del self.employees[name]['advance_due_dates'][date]
                self.update_employee_lists()
                self.save_data()
                messagebox.showinfo("تم", "تم حذف السلفة")
                self.clear_advance_fields()
        elif name in self.other_employees and date in self.other_employees[name].get('advances', {}):
            if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف سلفة {name} بتاريخ {date}؟"):
                del self.other_employees[name]['advances'][date]
                if 'advance_due_dates' in self.other_employees[name]:
                    del self.other_employees[name]['advance_due_dates'][date]
                self.update_employee_lists()
                self.save_data()
                messagebox.showinfo("تم", "تم حذف السلفة")
                self.clear_advance_fields()
        else:
            messagebox.showerror("خطأ", "السلفة غير موجودة")
    
    def clear_advance_fields(self):
        self.adv_amount.delete(0, tk.END)

    def create_reports_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="التقارير")
        
        # Report Controls Frame
        controls_frame = ttk.LabelFrame(tab, text="إعداد التقرير", padding=10)
        controls_frame.grid(row=0, column=0, padx=10, pady=10, sticky="ew")
        
        # Date range for reports
        ttk.Label(controls_frame, text="من تاريخ:").grid(row=0, column=0, padx=5, pady=5)
        self.report_from_date = ttk.Entry(controls_frame)
        self.report_from_date.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(controls_frame, text="إلى تاريخ:").grid(row=0, column=2, padx=5, pady=5)
        self.report_to_date = ttk.Entry(controls_frame)
        self.report_to_date.grid(row=0, column=3, padx=5, pady=5)
        
        # Add button to reset dates
        ttk.Button(controls_frame, text="إعادة تعيين", command=self.reset_report_dates,
                  style='Edit.TButton').grid(row=0, column=4, padx=5, pady=5)
        
        ttk.Label(controls_frame, text="الموظف:").grid(row=1, column=0, padx=5, pady=5)
        self.report_employee = ttk.Combobox(controls_frame, values=[], state="readonly")
        self.report_employee.grid(row=1, column=1, padx=5, pady=5)
        
        # Month and year for monthly reports
        ttk.Label(controls_frame, text="الشهر:").grid(row=1, column=2, padx=5, pady=5)
        self.report_month = ttk.Combobox(controls_frame, values=list(range(1, 13)), state="readonly")
        self.report_month.current(self.current_month - 1)
        self.report_month.grid(row=1, column=3, padx=5, pady=5)
        
        ttk.Label(controls_frame, text="السنة:").grid(row=1, column=4, padx=5, pady=5)
        self.report_year = ttk.Combobox(controls_frame, values=list(range(2020, 2031)), state="readonly")
        self.report_year.current(datetime.datetime.now().year - 2020)
        self.report_year.grid(row=1, column=5, padx=5, pady=5)
        
        # Add button to set default dates for the selected month
        ttk.Button(controls_frame, text="تعيين تواريخ الشهر", command=self.set_month_dates,
                  style='Edit.TButton').grid(row=1, column=6, padx=5, pady=5)
        
        btn_frame = ttk.Frame(controls_frame)
        btn_frame.grid(row=2, column=0, columnspan=7, pady=10)
        
        ttk.Button(btn_frame, text="عرض التقرير", command=self.generate_report,
                  style='Accent.TButton').pack(side='left', padx=5)
        ttk.Button(btn_frame, text="عرض تقرير الموظف", command=self.generate_employee_report,
                  style='Accent.TButton').pack(side='left', padx=5)
        ttk.Button(btn_frame, text="عرض تقرير الحضور", command=self.generate_attendance_report,
                  style='Accent.TButton').pack(side='left', padx=5)
        ttk.Button(btn_frame, text="تصدير لإكسل", command=self.export_to_excel,
                  style='Accent.TButton').pack(side='left', padx=5)
        ttk.Button(btn_frame, text="تصدير لوورد", command=self.export_to_word,
                  style='Export.TButton').pack(side='left', padx=5)
        
        # Report Display Frame
        display_frame = ttk.LabelFrame(tab, text="تقرير الرواتب", padding=10)
        display_frame.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")
        
        self.report_text = tk.Text(display_frame, wrap="none", font=('Arial', 12), bg='#ffffff')
        self.report_text.pack(fill="both", expand=True)
        
        scroll_y = ttk.Scrollbar(display_frame, orient="vertical", command=self.report_text.yview)
        scroll_y.pack(side="right", fill="y")
        scroll_x = ttk.Scrollbar(display_frame, orient="horizontal", command=self.report_text.xview)
        scroll_x.pack(side="bottom", fill="x")
        
        self.report_text.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        # Configure grid weights
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_rowconfigure(1, weight=1)
    
    def set_month_dates(self):
        """تعيين تواريخ بداية ونهاية الشهر المحدد"""
        month = int(self.report_month.get())
        year = int(self.report_year.get())
        
        # تحديد بداية ونهاية الشهر
        from_date = datetime.datetime(year, month, 1)
        to_date = datetime.datetime(year, month, 1) + datetime.timedelta(days=32)
        to_date = to_date.replace(day=1) - datetime.timedelta(days=1)
        
        self.report_from_date.delete(0, tk.END)
        self.report_from_date.insert(0, from_date.strftime('%Y-%m-%d'))
        
        self.report_to_date.delete(0, tk.END)
        self.report_to_date.insert(0, to_date.strftime('%Y-%m-%d'))
    
    def reset_report_dates(self):
        self.report_from_date.delete(0, tk.END)
        self.report_to_date.delete(0, tk.END)
    
    def calculate_salary_for_period(self, name, from_date=None, to_date=None):
        if name in self.employees:
            employee = self.employees[name]
            total_sessions = 0
            total_daily_bonus = 0
            base_rate = employee.get('current_rate', 0)
            
            # حساب الحضور في الفترة المحددة
            for date, record in employee.get('attendance', {}).items():
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                if from_date and date_obj < from_date:
                    continue
                if to_date and date_obj > to_date:
                    continue
                
                total_sessions += record.get('sessions', 0)
                total_daily_bonus += record.get('daily_bonus', 0)
            
            # حساب المكافآت والخصومات في الفترة المحددة
            performance_bonus = 0
            performance_sessions = 0
            monthly_bonus = 0
            deduction = 0
            advance_amount = 0
            
            # حساب بونص الأداء في الفترة المحددة
            for date, bonus_data in employee.get('performance_bonus', {}).items():
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                if from_date and date_obj < from_date:
                    continue
                if to_date and date_obj > to_date:
                    continue
                
                performance_sessions += bonus_data.get('sessions', 0)
                performance_bonus += bonus_data.get('amount', 0)
            
            # حساب البونص الشهري في الفترة المحددة
            for date, bonus in employee.get('monthly_bonuses', {}).items():
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                if from_date and date_obj < from_date:
                    continue
                if to_date and date_obj > to_date:
                    continue
                
                monthly_bonus += bonus
            
            # حساب الخصومات في الفترة المحددة
            for date, ded_data in employee.get('deductions', {}).items():
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                if from_date and date_obj < from_date:
                    continue
                if to_date and date_obj > to_date:
                    continue
                
                deduction += ded_data.get('amount', 0)
            
            # حساب السلف في الفترة المحددة
            for date, amount in employee.get('advances', {}).items():
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                if from_date and date_obj < from_date:
                    continue
                if to_date and date_obj > to_date:
                    continue
                
                due_date = employee.get('advance_due_dates', {}).get(date, {})
                due_month = due_date.get('month', 0)
                due_year = due_date.get('year', 0)
                
                if due_month and due_year:
                    due_date_obj = datetime.datetime(due_year, due_month, 1)
                    if from_date and due_date_obj < from_date:
                        continue
                    if to_date and due_date_obj > to_date:
                        continue
                
                advance_amount += amount
            
            # حساب سعر الحصة (آخر سعر تم تحديده في الفترة)
            current_rate = base_rate
            for date, rate in sorted(employee.get('monthly_rates', {}).items(), reverse=True):
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                if from_date and date_obj < from_date:
                    continue
                if to_date and date_obj > to_date:
                    continue
                
                current_rate = rate
                break
            
            # حساب راتب الحصص (عدد الحصص × سعر الحصة)
            sessions_salary = total_sessions * current_rate
            
            salary = sessions_salary + performance_bonus + total_daily_bonus + monthly_bonus - deduction - advance_amount
            
            return {
                'name': name,
                'type': 'بحصص',
                'base_rate': base_rate,
                'current_rate': current_rate,
                'sessions': total_sessions,
                'sessions_salary': sessions_salary,
                'performance_sessions': performance_sessions,
                'performance_bonus': performance_bonus,
                'daily_bonus': total_daily_bonus,
                'monthly_bonus': monthly_bonus,
                'deduction': deduction,
                'advance': advance_amount,
                'salary': salary,
                'from_date': from_date.strftime('%Y-%m-%d') if from_date else '',
                'to_date': to_date.strftime('%Y-%m-%d') if to_date else ''
            }
        elif name in self.other_employees:
            employee = self.other_employees[name]
            base_salary = employee.get('monthly_salary', 0)
            monthly_salary = 0
            monthly_bonus = 0
            deduction = 0
            advance_amount = 0
            
            # حساب الرواتب الشهرية في الفترة المحددة
            for month_year, salary in employee.get('monthly_salaries', {}).items():
                month, year = map(int, month_year.split('_'))
                month_start = datetime.datetime(year, month, 1)
                month_end = datetime.datetime(year, month, 1) + datetime.timedelta(days=32)
                month_end = month_end.replace(day=1) - datetime.timedelta(days=1)
                
                if from_date and month_end < from_date:
                    continue
                if to_date and month_start > to_date:
                    continue
                
                monthly_salary += salary
            
            # إذا لم يكن هناك رواتب محددة، نستخدم الراتب الأساسي
            if monthly_salary == 0 and base_salary > 0:
                # حساب عدد الأشهر في الفترة
                if from_date and to_date:
                    months = (to_date.year - from_date.year) * 12 + (to_date.month - from_date.month) + 1
                    monthly_salary = base_salary * months
            
            # حساب البونص الشهري في الفترة المحددة
            for date, bonus in employee.get('monthly_bonuses', {}).items():
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                if from_date and date_obj < from_date:
                    continue
                if to_date and date_obj > to_date:
                    continue
                
                monthly_bonus += bonus
            
            # حساب الخصومات في الفترة المحددة
            for date, ded_data in employee.get('deductions', {}).items():
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                if from_date and date_obj < from_date:
                    continue
                if to_date and date_obj > to_date:
                    continue
                
                deduction += ded_data.get('amount', 0)
            
            # حساب السلف في الفترة المحددة
            for date, amount in employee.get('advances', {}).items():
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                if from_date and date_obj < from_date:
                    continue
                if to_date and date_obj > to_date:
                    continue
                
                due_date = employee.get('advance_due_dates', {}).get(date, {})
                due_month = due_date.get('month', 0)
                due_year = due_date.get('year', 0)
                
                if due_month and due_year:
                    due_date_obj = datetime.datetime(due_year, due_month, 1)
                    if from_date and due_date_obj < from_date:
                        continue
                    if to_date and due_date_obj > to_date:
                        continue
                
                advance_amount += amount
            
            salary = monthly_salary + monthly_bonus - deduction - advance_amount
            
            return {
                'name': name,
                'type': 'راتب ثابت',
                'base_salary': base_salary,
                'monthly_salary': monthly_salary,
                'monthly_bonus': monthly_bonus,
                'deduction': deduction,
                'advance': advance_amount,
                'salary': salary,
                'from_date': from_date.strftime('%Y-%m-%d') if from_date else '',
                'to_date': to_date.strftime('%Y-%m-%d') if to_date else ''
            }
        else:
            return None
    
    def generate_report(self):
        from_date_str = self.report_from_date.get()
        to_date_str = self.report_to_date.get()
        
        from_date = None
        to_date = None
        
        if not from_date_str and not to_date_str:
            # إذا لم يتم تحديد تاريخ، نستخدم الشهر المحدد فقط
            month = int(self.report_month.get())
            year = int(self.report_year.get())
            from_date = datetime.datetime(year, month, 1)
            to_date = datetime.datetime(year, month, 1) + datetime.timedelta(days=32)
            to_date = to_date.replace(day=1) - datetime.timedelta(days=1)
        else:
            if from_date_str:
                try:
                    from_date = datetime.datetime.strptime(from_date_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("خطأ", "صيغة تاريخ البداية غير صحيحة")
                    return
            
            if to_date_str:
                try:
                    to_date = datetime.datetime.strptime(to_date_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("خطأ", "صيغة تاريخ النهاية غير صحيحة")
                    return
        
        if from_date and to_date and from_date > to_date:
            messagebox.showerror("خطأ", "تاريخ البداية يجب أن يكون قبل تاريخ النهاية")
            return
        
        reports = []
        total_salaries = 0
        
        # Regular employees
        for name in self.employees:
            report = self.calculate_salary_for_period(name, from_date, to_date)
            if report:
                reports.append(report)
                total_salaries += report['salary']
        
        # Other employees
        for name in self.other_employees:
            report = self.calculate_salary_for_period(name, from_date, to_date)
            if report:
                reports.append(report)
                total_salaries += report['salary']
        
        # Generate report text
        report_text = "\n" + "="*120 + "\n"
        report_text += " " * 40 + "تقرير الرواتب\n"
        
        if from_date and to_date:
            report_text += " " * 45 + f"للفترة من {from_date.strftime('%Y-%m-%d')} إلى {to_date.strftime('%Y-%m-%d')}\n"
        else:
            report_text += " " * 45 + f"لشهر {self.report_month.get()}/{self.report_year.get()}\n"
        
        report_text += "="*120 + "\n\n"
        
        # Regular employees details
        report_text += "الموظفون بحصص:\n"
        report_text += "-"*120 + "\n"
        
        headers = ["الاسم", "سعر الحصة", "الحصص العادية", "راتب الحصص", "حصص الأداء", "بونص الأداء", 
                  "البونص اليومي", "البونص الشهري", "الخصومات", "السلف", "صافي الراتب"]
        data = []
        
        for emp in reports:
            if emp['type'] == 'بحصص':
                data.append([
                    emp['name'],
                    f"{emp['current_rate']:.2f}",
                    emp['sessions'],
                    f"{emp['sessions_salary']:.2f}",
                    emp['performance_sessions'],
                    f"{emp['performance_bonus']:.2f}",
                    f"{emp['daily_bonus']:.2f}",
                    f"{emp['monthly_bonus']:.2f}",
                    f"{emp['deduction']:.2f}",
                    f"{emp['advance']:.2f}",
                    f"{emp['salary']:.2f}"
                ])
        
        if data:
            report_text += tabulate(data, headers=headers, tablefmt="pretty") + "\n\n"
        else:
            report_text += "لا يوجد موظفون بحصص\n\n"
        
        # Other employees details
        report_text += "الموظفون براتب ثابت:\n"
        report_text += "-"*120 + "\n"
        
        headers = ["الاسم", "الراتب الأساسي", "الراتب الشهري", "البونص الشهري", "الخصومات", "السلف", "صافي الراتب"]
        data = []
        
        for emp in reports:
            if emp['type'] == 'راتب ثابت':
                data.append([
                    emp['name'],
                    f"{emp['base_salary']:.2f}",
                    f"{emp['monthly_salary']:.2f}",
                    f"{emp['monthly_bonus']:.2f}",
                    f"{emp['deduction']:.2f}",
                    f"{emp['advance']:.2f}",
                    f"{emp['salary']:.2f}"
                ])
        
        if data:
            report_text += tabulate(data, headers=headers, tablefmt="pretty") + "\n\n"
        else:
            report_text += "لا يوجد موظفون براتب ثابت\n\n"
        
        # Summary
        report_text += "="*120 + "\n"
        report_text += " " * 40 + f"إجمالي رواتب جميع الموظفين: {total_salaries:.2f}\n"
        report_text += "="*120 + "\n"
        
        # Display report
        self.report_text.delete(1.0, "end")
        self.report_text.insert("end", report_text)
    
    def generate_employee_report(self):
        name = self.report_employee.get()
        if not name:
            messagebox.showerror("خطأ", "الرجاء تحديد الموظف")
            return
        
        from_date_str = self.report_from_date.get()
        to_date_str = self.report_to_date.get()
        
        from_date = None
        to_date = None
        
        if not from_date_str and not to_date_str:
            # إذا لم يتم تحديد تاريخ، نستخدم الشهر المحدد فقط
            month = int(self.report_month.get())
            year = int(self.report_year.get())
            from_date = datetime.datetime(year, month, 1)
            to_date = datetime.datetime(year, month, 1) + datetime.timedelta(days=32)
            to_date = to_date.replace(day=1) - datetime.timedelta(days=1)
        else:
            if from_date_str:
                try:
                    from_date = datetime.datetime.strptime(from_date_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("خطأ", "صيغة تاريخ البداية غير صحيحة")
                    return
            
            if to_date_str:
                try:
                    to_date = datetime.datetime.strptime(to_date_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("خطأ", "صيغة تاريخ النهاية غير صحيحة")
                    return
        
        if from_date and to_date and from_date > to_date:
            messagebox.showerror("خطأ", "تاريخ البداية يجب أن يكون قبل تاريخ النهاية")
            return
        
        report = self.calculate_salary_for_period(name, from_date, to_date)
        
        if not report:
            messagebox.showerror("خطأ", "لا يوجد بيانات لهذا الموظف")
            return
        
        # Generate report text
        report_text = "\n" + "="*100 + "\n"
        report_text += " " * 35 + "تقرير الموظف\n"
        
        if from_date and to_date:
            report_text += " " * 40 + f"للفترة من {from_date.strftime('%Y-%m-%d')} إلى {to_date.strftime('%Y-%m-%d')}\n"
        else:
            report_text += " " * 40 + f"لشهر {self.report_month.get()}/{self.report_year.get()}\n"
        
        report_text += "="*100 + "\n\n"
        
        report_text += f"الاسم: {report['name']}\n"
        report_text += f"النوع: {report['type']}\n"
        
        if report['type'] == 'بحصص':
            report_text += f"سعر الحصة الأساسي: {report['base_rate']:.2f}\n"
            report_text += f"سعر الحصة الحالي: {report['current_rate']:.2f}\n"
            report_text += f"عدد الحصص العادية: {report['sessions']}\n"
            report_text += f"راتب الحصص: {report['sessions_salary']:.2f}\n"
            report_text += f"عدد حصص الأداء: {report['performance_sessions']}\n"
            report_text += f"بونص الأداء: {report['performance_bonus']:.2f}\n"
            report_text += f"إجمالي البونص اليومي: {report['daily_bonus']:.2f}\n"
            report_text += f"البونص الشهري: {report['monthly_bonus']:.2f}\n"
        else:
            report_text += f"الراتب الأساسي: {report['base_salary']:.2f}\n"
            report_text += f"الراتب الشهري: {report['monthly_salary']:.2f}\n"
            report_text += f"البونص الشهري: {report['monthly_bonus']:.2f}\n"
        
        report_text += f"الخصومات: {report['deduction']:.2f}\n"
        report_text += f"السلف المستحقة: {report['advance']:.2f}\n"
        report_text += "-"*100 + "\n"
        report_text += f"صافي الراتب: {report['salary']:.2f}\n"
        report_text += "="*100 + "\n"
        
        # Display report
        self.report_text.delete(1.0, "end")
        self.report_text.insert("end", report_text)
    
    def generate_attendance_report(self):
        name = self.report_employee.get()
        if not name:
            messagebox.showerror("خطأ", "الرجاء تحديد الموظف")
            return
        
        if name not in self.employees:
            messagebox.showerror("خطأ", "هذا الموظف ليس من الموظفين بحصص")
            return
        
        from_date_str = self.report_from_date.get()
        to_date_str = self.report_to_date.get()
        
        from_date = None
        to_date = None
        
        if not from_date_str and not to_date_str:
            # إذا لم يتم تحديد تاريخ، نستخدم الشهر المحدد فقط
            month = int(self.report_month.get())
            year = int(self.report_year.get())
            from_date = datetime.datetime(year, month, 1)
            to_date = datetime.datetime(year, month, 1) + datetime.timedelta(days=32)
            to_date = to_date.replace(day=1) - datetime.timedelta(days=1)
        else:
            if from_date_str:
                try:
                    from_date = datetime.datetime.strptime(from_date_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("خطأ", "صيغة تاريخ البداية غير صحيحة")
                    return
            
            if to_date_str:
                try:
                    to_date = datetime.datetime.strptime(to_date_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("خطأ", "صيغة تاريخ النهاية غير صحيحة")
                    return
        
        if from_date and to_date and from_date > to_date:
            messagebox.showerror("خطأ", "تاريخ البداية يجب أن يكون قبل تاريخ النهاية")
            return
        
        attendance_records = []
        for date, record in self.employees[name].get('attendance', {}).items():
            date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
            
            # Apply date range filter if specified
            if from_date and date_obj < from_date:
                continue
            if to_date and date_obj > to_date:
                continue
            
            day_name = ["الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت", "الأحد"][date_obj.weekday()]
            
            attendance_records.append({
                'date': date,
                'day': day_name,
                'sessions': record.get('sessions', 0),
                'daily_bonus': record.get('daily_bonus', 0)
            })
        
        if not attendance_records:
            messagebox.showinfo("السجل", "لا يوجد سجل حضور لهذا الموظف في الفترة المحددة")
            return
        
        # ترتيب التواريخ من الأقدم إلى الأحدث
        attendance_records.sort(key=lambda x: x['date'])
        
        # إنشاء DataFrame للتصدير
        df = pd.DataFrame(attendance_records)
        
        # عرض التقرير في النافذة
        report_text = f"\nتقرير حضور الموظف {name}\n"
        
        if from_date and to_date:
            report_text += f"للفترة من {from_date.strftime('%Y-%m-%d')} إلى {to_date.strftime('%Y-%m-%d')}\n"
        else:
            report_text += f"لشهر {self.report_month.get()}/{self.report_year.get()}\n"
        
        report_text += "-"*80 + "\n"
        report_text += tabulate(df, headers=['التاريخ', 'اليوم', 'الحصص', 'البونص اليومي'], tablefmt="pretty")
        
        # Calculate totals
        total_sessions = sum(record['sessions'] for record in attendance_records)
        total_bonus = sum(record['daily_bonus'] for record in attendance_records)
        
        report_text += "\n\n" + "-"*80 + "\n"
        report_text += f"الإجمالي: الحصص: {total_sessions} | البونص اليومي: {total_bonus:.2f}\n"
        report_text += "="*80 + "\n"
        
        self.report_text.delete(1.0, "end")
        self.report_text.insert("end", report_text)
        
        # تصدير إلى Excel
        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
            title="حفظ تقرير الحضور كملف إكسل"
        )
        
        if file_path:
            try:
                df.to_excel(file_path, index=False)
                messagebox.showinfo("تم", f"تم تصدير تقرير الحضور إلى {file_path}")
            except Exception as e:
                messagebox.showerror("خطأ", f"فشل التصدير: {str(e)}")
    
    def export_to_excel(self):
        from_date_str = self.report_from_date.get()
        to_date_str = self.report_to_date.get()
        
        from_date = None
        to_date = None
        
        if not from_date_str and not to_date_str:
            # إذا لم يتم تحديد تاريخ، نستخدم الشهر المحدد فقط
            month = int(self.report_month.get())
            year = int(self.report_year.get())
            from_date = datetime.datetime(year, month, 1)
            to_date = datetime.datetime(year, month, 1) + datetime.timedelta(days=32)
            to_date = to_date.replace(day=1) - datetime.timedelta(days=1)
        else:
            if from_date_str:
                try:
                    from_date = datetime.datetime.strptime(from_date_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("خطأ", "صيغة تاريخ البداية غير صحيحة")
                    return
            
            if to_date_str:
                try:
                    to_date = datetime.datetime.strptime(to_date_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("خطأ", "صيغة تاريخ النهاية غير صحيحة")
                    return
        
        if from_date and to_date and from_date > to_date:
            messagebox.showerror("خطأ", "تاريخ البداية يجب أن يكون قبل تاريخ النهاية")
            return
        
        reports = []
        total_salaries = 0
        
        # Regular employees
        for name in self.employees:
            report = self.calculate_salary_for_period(name, from_date, to_date)
            if report:
                reports.append(report)
                total_salaries += report['salary']
        
        # Other employees
        for name in self.other_employees:
            report = self.calculate_salary_for_period(name, from_date, to_date)
            if report:
                reports.append(report)
                total_salaries += report['salary']
        
        # Prepare data for Excel
        data = []
        
        for emp in reports:
            if emp['type'] == 'بحصص':
                data.append({
                    'الاسم': emp['name'],
                    'النوع': 'بحصص',
                    'سعر الحصة الأساسي': f"{emp['base_rate']:.2f}",
                    'سعر الحصة الحالي': f"{emp['current_rate']:.2f}",
                    'الحصص العادية': emp['sessions'],
                    'راتب الحصص': f"{emp['sessions_salary']:.2f}",
                    'حصص الأداء': emp['performance_sessions'],
                    'بونص الأداء': f"{emp['performance_bonus']:.2f}",
                    'البونص اليومي': f"{emp['daily_bonus']:.2f}",
                    'البونص الشهري': f"{emp['monthly_bonus']:.2f}",
                    'الخصومات': f"{emp['deduction']:.2f}",
                    'السلف': f"{emp['advance']:.2f}",
                    'صافي الراتب': f"{emp['salary']:.2f}"
                })
            else:
                data.append({
                    'الاسم': emp['name'],
                    'النوع': 'راتب ثابت',
                    'الراتب الأساسي': f"{emp['base_salary']:.2f}",
                    'الراتب الشهري': f"{emp['monthly_salary']:.2f}",
                    'البونص الشهري': f"{emp['monthly_bonus']:.2f}",
                    'الخصومات': f"{emp['deduction']:.2f}",
                    'السلف': f"{emp['advance']:.2f}",
                    'صافي الراتب': f"{emp['salary']:.2f}"
                })
        
        # Add total row
        data.append({
            'الاسم': 'الإجمالي',
            'النوع': '',
            'سعر الحصة الأساسي': '',
            'سعر الحصة الحالي': '',
            'الحصص العادية': '',
            'راتب الحصص': '',
            'حصص الأداء': '',
            'بونص الأداء': '',
            'البونص اليومي': '',
            'البونص الشهري': '',
            'الراتب الأساسي': '',
            'الراتب الشهري': '',
            'الخصومات': '',
            'السلف': '',
            'صافي الراتب': f"{total_salaries:.2f}"
        })
        
        df = pd.DataFrame(data)
        
        # Ask for save location
        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
            title="حفظ التقرير كملف إكسل"
        )
        
        if file_path:
            try:
                df.to_excel(file_path, index=False)
                messagebox.showinfo("تم", f"تم تصدير التقرير إلى {file_path}")
            except Exception as e:
                messagebox.showerror("خطأ", f"فشل التصدير: {str(e)}")
    
    def export_to_word(self):
        name = self.report_employee.get()
        if not name:
            messagebox.showerror("خطأ", "الرجاء تحديد الموظف")
            return
        
        from_date_str = self.report_from_date.get()
        to_date_str = self.report_to_date.get()
        
        from_date = None
        to_date = None
        
        if not from_date_str and not to_date_str:
            # إذا لم يتم تحديد تاريخ، نستخدم الشهر المحدد فقط
            month = int(self.report_month.get())
            year = int(self.report_year.get())
            from_date = datetime.datetime(year, month, 1)
            to_date = datetime.datetime(year, month, 1) + datetime.timedelta(days=32)
            to_date = to_date.replace(day=1) - datetime.timedelta(days=1)
        else:
            if from_date_str:
                try:
                    from_date = datetime.datetime.strptime(from_date_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("خطأ", "صيغة تاريخ البداية غير صحيحة")
                    return
            
            if to_date_str:
                try:
                    to_date = datetime.datetime.strptime(to_date_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror("خطأ", "صيغة تاريخ النهاية غير صحيحة")
                    return
        
        if from_date and to_date and from_date > to_date:
            messagebox.showerror("خطأ", "تاريخ البداية يجب أن يكون قبل تاريخ النهاية")
            return
        
        report = self.calculate_salary_for_period(name, from_date, to_date)
        
        if not report:
            messagebox.showerror("خطأ", "لا يوجد بيانات لهذا الموظف")
            return
        
        # Create a new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('تقرير الموظف', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle_text = f"الموظف: {report['name']}"
        if from_date and to_date:
            subtitle_text += f" - الفترة من {from_date.strftime('%Y-%m-%d')} إلى {to_date.strftime('%Y-%m-%d')}"
        else:
            subtitle_text += f" - لشهر {self.report_month.get()}/{self.report_year.get()}"
        
        subtitle = doc.add_heading(subtitle_text, level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a table for the report
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table width to 100% of page
        table.autofit = False
        table.allow_autofit = False
        table.width = Inches(6)
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'القيمة'
        hdr_cells[1].text = 'البند'
        
        # Add shading to header
        shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        hdr_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
        
        # Add employee details to the table
        def add_row(label, value):
            row_cells = table.add_row().cells
            row_cells[0].text = str(value)
            row_cells[1].text = label
            # Set cell width
            row_cells[0].width = Inches(4)
            row_cells[1].width = Inches(2)
        
        add_row('الاسم', report['name'])
        add_row('النوع', report['type'])
        
        if report['type'] == 'بحصص':
            add_row('سعر الحصة الأساسي', f"{report['base_rate']:.2f}")
            add_row('سعر الحصة الحالي', f"{report['current_rate']:.2f}")
            add_row('عدد الحصص العادية', report['sessions'])
            add_row('راتب الحصص', f"{report['sessions_salary']:.2f}")
            add_row('عدد حصص الأداء', report['performance_sessions'])
            add_row('بونص الأداء', f"{report['performance_bonus']:.2f}")
            add_row('إجمالي البونص اليومي', f"{report['daily_bonus']:.2f}")
            add_row('البونص الشهري', f"{report['monthly_bonus']:.2f}")
        else:
            add_row('الراتب الأساسي', f"{report['base_salary']:.2f}")
            add_row('الراتب الشهري', f"{report['monthly_salary']:.2f}")
            add_row('البونص الشهري', f"{report['monthly_bonus']:.2f}")
        
        add_row('الخصومات', f"{report['deduction']:.2f}")
        add_row('السلف المستحقة', f"{report['advance']:.2f}")
        
        # Add total salary row with bold font and different background color
        total_row = table.add_row().cells
        total_row[0].text = f"{report['salary']:.2f}"
        total_row[1].text = 'صافي الراتب'
        
        # Apply bold font
        for cell in total_row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Apply background color
        shading_elm = parse_xml(r'<w:shd {} w:fill="B4C6E7"/>'.format(nsdecls('w')))
        total_row[0]._tc.get_or_add_tcPr().append(shading_elm)
        total_row[1]._tc.get_or_add_tcPr().append(shading_elm)
        
        # Add date and signature
        doc.add_paragraph("\n")
        doc.add_paragraph(f"تاريخ التقرير: {datetime.date.today().strftime('%Y-%m-%d')}")
        doc.add_paragraph("توقيع المدير: ________________")
        
        # Save the document
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
            title="حفظ التقرير كملف وورد"
        )
        
        if file_path:
            try:
                doc.save(file_path)
                messagebox.showinfo("تم", f"تم تصدير التقرير إلى {file_path}")
            except Exception as e:
                messagebox.showerror("خطأ", f"فشل التصدير: {str(e)}")

    def update_employee_lists(self):
        # Update comboboxes
        all_employees = list(self.employees.keys()) + list(self.other_employees.keys())
        self.bonus_employee['values'] = list(self.employees.keys())
        self.ded_employee['values'] = all_employees
        self.report_employee['values'] = all_employees
        self.adv_employee['values'] = all_employees
        
        # Update employee treeview
        self.employee_tree.delete(*self.employee_tree.get_children())
        
        for name, data in self.employees.items():
            self.employee_tree.insert("", "end", values=(
                name,
                "بحصص",
                f"{data.get('current_rate', 0):.2f}",
                "حذف"
            ), tags=('editable',))
        
        for name, data in self.other_employees.items():
            self.employee_tree.insert("", "end", values=(
                name,
                "راتب ثابت",
                f"{data.get('monthly_salary', 0):.2f}",
                "حذف"
            ), tags=('editable',))
        
        # Update attendance treeview with day names
        self.attendance_tree.delete(*self.attendance_tree.get_children())
        
        # ترتيب الموظفين حسب الاسم
        sorted_employees = sorted(self.employees.items(), key=lambda x: x[0])
        
        for name, data in sorted_employees:
            # ترتيب سجلات الحضور حسب التاريخ
            sorted_attendance = sorted(data.get('attendance', {}).items(), key=lambda x: x[0])
            
            for date, record in sorted_attendance:
                date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                day_name = ["الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت", "الأحد"][date_obj.weekday()]
                
                self.attendance_tree.insert("", "end", values=(
                    name,
                    date,
                    day_name,
                    record.get('sessions', 0),
                    f"{record.get('daily_bonus', 0):.2f}",
                    "تعديل",
                    "حذف"
                ), tags=('editable',))
        
        # Update bonus treeview
        self.bonus_tree.delete(*self.bonus_tree.get_children())
        
        for name, data in self.employees.items():
            # Add performance bonuses
            for date, bonus_data in data.get('performance_bonus', {}).items():
                sessions = bonus_data.get('sessions', 0)
                rate = bonus_data.get('rate', 0)
                amount = bonus_data.get('amount', 0)
                
                self.bonus_tree.insert("", "end", values=(
                    name,
                    date,
                    "بونص الأداء",
                    f"{sessions} حصة × {rate:.2f}",
                    f"{amount:.2f}",
                    "تعديل",
                    "حذف"
                ))
            
            # Add monthly bonuses
            for date, bonus in data.get('monthly_bonuses', {}).items():
                if bonus > 0:
                    self.bonus_tree.insert("", "end", values=(
                        name,
                        date,
                        "البونص الشهري",
                        "-",
                        f"{bonus:.2f}",
                        "تعديل",
                        "حذف"
                    ))
        
        # Update advance treeview
        self.advance_tree.delete(*self.advance_tree.get_children())
        
        for name, data in self.employees.items():
            for date, amount in data.get('advances', {}).items():
                due_month = data.get('advance_due_dates', {}).get(date, {}).get('month', '')
                due_year = data.get('advance_due_dates', {}).get(date, {}).get('year', '')
                self.advance_tree.insert("", "end", values=(
                    name,
                    date,
                    f"{amount:.2f}",
                    due_month,
                    due_year,
                    "حذف"
                ), tags=('editable',))
        
        for name, data in self.other_employees.items():
            for date, amount in data.get('advances', {}).items():
                due_month = data.get('advance_due_dates', {}).get(date, {}).get('month', '')
                due_year = data.get('advance_due_dates', {}).get(date, {}).get('year', '')
                self.advance_tree.insert("", "end", values=(
                    name,
                    date,
                    f"{amount:.2f}",
                    due_month,
                    due_year,
                    "حذف"
                ), tags=('editable',))
    
    def save_regular_employee(self):
        name = self.reg_name.get()
        rate = self.reg_rate.get()
        
        if not name or not rate:
            messagebox.showerror("خطأ", "الاسم وسعر الحصة مطلوبان")
            return
        
        if name in self.employees or name in self.other_employees:
            messagebox.showerror("خطأ", "اسم الموظف موجود بالفعل")
            return
        
        try:
            rate = float(rate)
        except ValueError:
            messagebox.showerror("خطأ", "الرجاء إدخال أرقام صحيحة")
            return
        
        self.employees[name] = {
            'current_rate': rate,
            'attendance': {},
            'performance_bonus': {},
            'monthly_bonuses': {},
            'deductions': {},
            'advances': {},
            'advance_due_dates': {},
            'monthly_rates': {}
        }
        
        self.update_employee_lists()
        self.save_data()
        messagebox.showinfo("تم", f"تم حفظ بيانات الموظف {name}")
        self.clear_regular_employee_fields()
    
    def delete_regular_employee(self):
        name = self.reg_name.get()
        
        if not name:
            messagebox.showerror("خطأ", "الرجاء إدخال اسم الموظف")
            return
        
        if name not in self.employees:
            messagebox.showerror("خطأ", "الموظف غير موجود")
            return
        
        if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف الموظف {name}؟"):
            del self.employees[name]
            self.update_employee_lists()
            self.save_data()
            messagebox.showinfo("تم", f"تم حذف الموظف {name}")
            self.clear_regular_employee_fields()
    
    def clear_regular_employee_fields(self):
        self.reg_name.delete(0, tk.END)
        self.reg_rate.delete(0, tk.END)
    
    def save_other_employee(self):
        name = self.other_name.get()
        salary = self.other_salary.get()
        
        if not name or not salary:
            messagebox.showerror("خطأ", "الاسم والراتب الأساسي مطلوبان")
            return
        
        if name in self.employees or name in self.other_employees:
            messagebox.showerror("خطأ", "اسم الموظف موجود بالفعل")
            return
        
        try:
            salary = float(salary)
        except ValueError:
            messagebox.showerror("خطأ", "الرجاء إدخال أرقام صحيحة")
            return
        
        self.other_employees[name] = {
            'monthly_salary': salary,
            'monthly_salaries': {},
            'monthly_bonuses': {},
            'deductions': {},
            'advances': {},
            'advance_due_dates': {}
        }
        
        self.update_employee_lists()
        self.save_data()
        messagebox.showinfo("تم", f"تم حفظ بيانات الموظف {name}")
        self.clear_other_employee_fields()
    
    def update_salary(self):
        name = self.other_name.get()
        month = self.other_month.get()
        year = self.other_year.get()
        monthly_salary = self.other_monthly_salary.get()
        
        if not name or not month or not year or not monthly_salary:
            messagebox.showerror("خطأ", "جميع الحقول مطلوبة")
            return
        
        try:
            month = int(month)
            year = int(year)
            monthly_salary = float(monthly_salary)
        except ValueError:
            messagebox.showerror("خطأ", "الرجاء إدخال أرقام صحيحة")
            return
        
        if name not in self.other_employees:
            messagebox.showerror("خطأ", "الموظف غير موجود")
            return
        
        month_year = f"{month}_{year}"
        
        if 'monthly_salaries' not in self.other_employees[name]:
            self.other_employees[name]['monthly_salaries'] = {}
        
        self.other_employees[name]['monthly_salaries'][month_year] = monthly_salary
        
        self.save_data()
        messagebox.showinfo("تم", f"تم تحديث الراتب الشهري لـ {name} لشهر {month}/{year}")
        self.clear_other_employee_fields()
    
    def delete_other_employee(self):
        name = self.other_name.get()
        
        if not name:
            messagebox.showerror("خطأ", "الرجاء إدخال اسم الموظف")
            return
        
        if name not in self.other_employees:
            messagebox.showerror("خطأ", "الموظف غير موجود")
            return
        
        if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف الموظف {name}؟"):
            del self.other_employees[name]
            self.update_employee_lists()
            self.save_data()
            messagebox.showinfo("تم", f"تم حذف الموظف {name}")
            self.clear_other_employee_fields()
    
    def clear_other_employee_fields(self):
        self.other_monthly_salary.delete(0, tk.END)

# Run the application
if __name__ == "__main__":
    root = tk.Tk()
    app = EnhancedEmployeeSystem(root)
    
    def on_closing():
        app.save_data()
        root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", on_closing)
    root.mainloop() 